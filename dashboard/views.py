from core.models import UserProfile, UserAccount
from django.shortcuts import get_object_or_404, render
from django.contrib.auth.models import User
from django.http import HttpResponse, JsonResponse
from .models import UserAccount, MasterModule, RequestModule, Resource
from core.models import UserProfile
from django.shortcuts import render, HttpResponse
from django.contrib.staticfiles import finders
from django.template.loader import get_template
from django.conf import settings
from django.views.decorators.csrf import csrf_exempt
from django.http import JsonResponse
from django.shortcuts import render, redirect
from django.contrib.auth.decorators import login_required
from .models import *
import base64
from django.shortcuts import get_object_or_404
from .models import MasterModule, Resource
from django.core import serializers
from collections import OrderedDict
from django.contrib.auth.decorators import login_required
from django.db.models import Q
from django.core.serializers.json import DjangoJSONEncoder
from django.core.files.storage import default_storage
from django.shortcuts import render, HttpResponse, redirect
from docx2pdf import convert
from django.shortcuts import render
from django.utils import timezone
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.http import JsonResponse, HttpResponseRedirect
from .models import RequestModule, ReviewModule
import tempfile
from django.core.paginator import Paginator, PageNotAnInteger, EmptyPage
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.platypus import Paragraph
from docx.text.run import Run
from reportlab.lib.styles import getSampleStyleSheet
from docx.oxml import OxmlElement
import os
from PyPDF2 import PdfReader, PdfWriter, PageObject
from PyPDF2.generic import AnnotationBuilder
import docx
from collections import OrderedDict
import re
import json
from django.core import serializers
from django.forms.models import model_to_dict
from django.shortcuts import render, get_object_or_404
from django.template.loader import render_to_string
from weasyprint import HTML
from rest_framework.response import Response
from .serializers import RequestSerializer
from rest_framework.decorators import api_view
from datetime import datetime
from .automations import create_custom_review
from .models import UserAccount, MasterModule, RequestModule, Resource
from django.contrib.auth.decorators import login_required 
from django.shortcuts import get_object_or_404, render
from docx.enum.text import WD_COLOR_INDEX, WD_PARAGRAPH_ALIGNMENT
from docx import Document
import filetype
import json

# ok
@login_required(login_url='userLogin')
def dashboard(request):
    total_requests = RequestModule.objects.filter(username=request.user).count()
    total_reviews = ReviewModule.objects.filter(username=request.user).count()
    context = {
        'total_requests': total_requests,
        'total_reviews': total_reviews
    }
    return render(request, 'dashboard/dashboard.html', {'Title': 'Dashboard', 'context': context})

# ok
@login_required(login_url='userLogin')
def edit_profile(request):
    user_profile = UserProfile.objects.get(userAccount=request.user)
    if request.method == 'POST':
        fullname = request.POST.get('fullname')
        email = request.POST.get('email')
        area = request.POST.get('area')
        domain = request.POST.get('domain')
        call = request.POST.get('call')
        Whatsapp = request.POST.get('Whatsapp')
        state = request.POST.get('state')
        city = request.POST.get('city')
        Address = request.POST.get('Address')
        about = request.POST.get('about')
        current_user = UserAccount.objects.get(username=request.user)
        current_user.full_name = fullname
        current_user.email = email
        current_user.save()
        user_profile.area = area
        user_profile.domain = domain
        user_profile.phone_no = call
        user_profile.whatsApp_no = Whatsapp
        user_profile.state = state
        user_profile.city = city
        user_profile.address = Address
        user_profile.about = about
        user_profile.save()
        return redirect('edit_profile')

    return render(request, 'dashboard/profile.html', {'Title': 'Profile', 'data': user_profile})


''' DO NOT UPDATE COMMENTS
************************************************
*   REQUEST FUNCTIONS                          *
************************************************
'''

# Create new request
@login_required(login_url='userLogin')
def create_request(request, id):
    if request.method == 'POST':
        # fetching required details
        user_account = UserAccount.objects.get(pk=request.user.id)
        request_name = request.POST.get('request')
        ModuleName = MasterModule.objects.get(id=id)
        json_field = request.POST.get('data')
        json_field = json.loads(json_field)

        # check if the user own same request name
        existing_request_per_user = RequestModule.objects.filter(username=request.user, request_name=request_name).exists()

        if not existing_request_per_user:
            request_save = RequestModule(username=user_account, request_name=request_name, ModuleName=ModuleName, json_field=json_field)
            request_save.save()
            return JsonResponse({'success': True, 'message': 'Request saved successfully'})
        else:
            return JsonResponse({'success': False, 'message': 'Request name already exist'})

    selectedModule = MasterModule.objects.get(id=id)
    json_data = selectedModule.json_field

    generalPoint = json_data['General Point']
    technicalPoint = json_data['Technical Point']
    resources = Resource.objects.get(id=1).json_field

    return render(request, "dashboard/create_request.html", {'Title': 'Request', 'selectedModule': selectedModule, 'standardmaster': 'master ', 'generalPoint': generalPoint, 'technicalPoint': technicalPoint, 'resources': resources})

# Show all request
@login_required(login_url='userLogin')
@api_view(['GET'])
def show_all_request(request):
    modulesList = MasterModule.objects.all()
    page = request.GET.get('page')
    fav = request.GET.get('fav')
    search_query = request.GET.get('search')
    date_query = request.GET.get('dateRange')

    all_request = RequestModule.objects.filter(
        username=request.user).order_by('-created_at')

    if search_query:
        all_request = all_request.filter(
            Q(request_name__icontains=search_query))

    if date_query:
        try:
            if date_query != None:
                start_date_str, end_date_str = date_query.split(" - ")
                formatted_start_date = timezone.datetime.strptime(
                    start_date_str, "%d/%m/%Y")
                formatted_end_date = timezone.datetime.strptime(
                    end_date_str, "%d/%m/%Y")
                all_request = all_request.filter(created_at__range=[
                                                 formatted_start_date, formatted_end_date + timezone.timedelta(days=1)])
        except ValueError:
            pass

    if fav == 'true':
        all_request = all_request.filter(is_favorite=True)

    items_per_page = 3
    paginator = Paginator(all_request, items_per_page)
    page = request.GET.get('page')

    if request.META.get('HTTP_X_REQUESTED_WITH') == 'XMLHttpRequest':
        try:
            all_request = paginator.page(page)
        except PageNotAnInteger:
            all_request = paginator.page(1)
        except EmptyPage:
            all_request = paginator.page(paginator.num_pages)

        serialized_request_list = RequestSerializer(all_request, many=True)

        all_pages = all_request.paginator.num_pages
        current_page = all_request.number

        has_previous = all_request.has_previous()
        if has_previous:
            previous_page_number = all_request.previous_page_number()
        else:
            previous_page_number = False

        has_next = all_request.has_next()
        if has_next:
            next_page_number = all_request.next_page_number()
        else:
            next_page_number = False

        data = {
            'requestList': serialized_request_list.data,
            'has_previous': has_previous,
            'has_next': has_next,
            'previous_page_number': previous_page_number,
            'next_page_number': next_page_number,
            'all_pages': all_pages,
            'current_page': current_page,
        }

        return Response(data)
        # return JsonResponse(data)

    return render(request, 'dashboard/show_all_request.html', {'Title': 'Request', "master": modulesList})


def save_note(request):
    if request.method == 'POST':
        noteContent = request.POST.get('noteContent', '')
        noteId = request.POST.get('noteId', '')

        request_module_instance = RequestModule.objects.get(
            username=request.user, id=noteId)
        if request_module_instance:
            request_module_instance.add_note = noteContent
            request_module_instance.save()

        return JsonResponse({'success': True})
    else:
        return JsonResponse({'success': False, 'error': 'Invalid request method'})


# Update request
@login_required(login_url='userLogin')
def update_request(request, id):
    if request.method == 'POST':
        request_name = request.POST.get('request')
        originalName = request.POST.get('originalName')
        json_field1 = request.POST.get('data')

        existing_request = RequestModule.objects.filter(
            username=request.user, request_name=request_name).exists()

        if request_name == originalName:
            request_save = RequestModule.objects.get(id=id)
            request_save.json_field = json.loads(json_field1)
            request_save.save()
            return JsonResponse({'success': True, 'message': 'Request Updated successfully'})
        elif request_name != originalName:
            request_save = RequestModule.objects.get(id=id)
            request_save.request_name = request_name
            request_save.json_field = json.loads(json_field1)
            request_save.save()
            return JsonResponse({'success': True, 'message': 'Request Update successfully'})
        else:
            return JsonResponse({'success': False, 'message': 'Request name already exist'})

    master = RequestModule.objects.get(id=id)
    json_data = master.json_field
    generalPoint = json_data['General Point']
    technicalPoint = json_data['Technical Point']
    selectedModule = MasterModule.objects.get(
        module_name=master.ModuleName.module_name)
    json_data_master = selectedModule.json_field
    generalPoint_master = json_data_master['General Point']
    technicalPoint_master = json_data_master['Technical Point']

    resources = Resource.objects.get(id=1).json_field
    return render(request, 'dashboard/update_request.html', {'Title': 'Request', 'standard': master, 'generalPoint': generalPoint, 'technicalPoint': technicalPoint, 'resources': resources, 'generalPoint_master': generalPoint_master, 'technicalPoint_master': technicalPoint_master, 'request_name': master.request_name})


# Delete request
@login_required(login_url='userLogin')
def delete_request(request, id):
    RequestModule.objects.get(id=id).delete()
    return redirect('show_all_request')


# Delete all requests
@csrf_exempt
def delete_all_request(request):
    if request.method == 'POST':
        selected_ids = request.POST.getlist('selected_ids[]')
        RequestModule.objects.filter(
            username=request.user, id__in=selected_ids).delete()
    return JsonResponse({'success': True, 'message': 'Request Delete successfully'})

# Generate request Pdf
@csrf_exempt
def create_request_pdf(request):
    if request.method == 'POST':
        id = request.POST.get('id')
        temp = RequestModule.objects.get(id=id)
        print(temp)
        data = temp.json_field
        print('==,', data)
        requestname = temp.request_name
        context = {
            "Technical": data["Technical Point"],
            "General": data["General Point"],
            'requestname': requestname,
            "clientname": request.user.username,
            "modulename": temp.ModuleName,
            "date": temp.created_at
        }
    # return render(request, 'dashboard/pdf.html', context)
        html_string = render_to_string('dashboard/pdf.html', context)
        pdf = HTML(string=html_string).write_pdf()
        pdf_base64 = base64.b64encode(pdf).decode('utf-8')

        # Include requestname in the response data
        response_data = {
            'requestname': requestname,
            'clientname': request.user.username,
            'pdf_data': pdf_base64,
        }
        return JsonResponse(response_data)


''' DO NOT UPDATE COMMENTS
************************************************
*   REVIEW FUNCTIONS                          *
************************************************
'''

# 1. Create review - Master (formally admin/default created)
@login_required(login_url='userLogin')
def create_master_review(request):
    if request.method == 'POST':
        docx_file = request.FILES['docx_file']
        request_name = request.POST.get('request_name',)
        checkpoints = MasterModule.objects.get(id=request_name)
        json_data = checkpoints.json_field
        general = json_data['General Point']
        technical = json_data['Technical Point']

        exclude_keys = ["General Point", "Technical Point"]
        keys = []

        def extract_keys_recursive(dictionary):
            for key, value in dictionary.items():
                keys.append(key)  # Include the key itself
                if isinstance(value, dict):
                    extract_keys_recursive(value)  # Recurse into nested dictionary
        extract_keys_recursive(general)

        
        def extract_keys_recursive(dictionary): 
            for key, value in dictionary.items():
                keys.append(key)  # Include the key itself 
                if isinstance(value, dict):
                    extract_keys_recursive(value)  # Recurse into nested dictionary
        extract_keys_recursive(technical)

        exclude_keys = ["General Point", "Font", "Page Layout","Technical Point","Additional Parameters","Abstract should have","Dataset","Keywords"]
        filtered_keys = [key for key in keys if key not in exclude_keys]
        word_file = docx_file

        # Define the keys and filtered_keys based on your requirements
        filtered_keys = ['Font Style', 'Font Size', 'Font Color', 'Background Color', 'Bold', 'Italic', 'Underline', 'Page Column', 'Spacing Before', 'Spacing After', 'Margin Top', 'Margin Left', 'Margin Right', 'Margin Bottom', 'Indent Right', 'Indent Left', 'Alignment', 'Words like we, I, our and you','Abstract Word Count','Document Type','Accepted Grammar Mistakes (Syntactically)','Accepted Plagiarism','Plag Tool','Problem Discussed','Objective','Novelty','Technique Used','Result (Mathematical Description)','Future Scope','Dataset Description','Dataset Source','Dataset Type','Keywords Heading','No. of keywords']

        # Function to extract font styles from the document
        def extract_font_styles(word_file):
            document = Document(word_file)
            font_styles = set()

            for paragraph in document.paragraphs:
                for run in paragraph.runs:
                    font_styles.add(run.font.name)

            return list(font_styles)

        # Function to extract font sizes from the document
        def extract_font_sizes(word_file):
            document = Document(word_file)
            font_sizes = set()

            for paragraph in document.paragraphs:
                for run in paragraph.runs:
                    if run.font.size is not None:
                        font_sizes.add(run.font.size.pt)

            return list(font_sizes)

        # Function to extract font colors from the document

        def extract_font_colors(word_file):
            document = Document(word_file)
            font_colors = set()

            for paragraph in document.paragraphs:
                for run in paragraph.runs:
                    font_colors.add(run.font.color.rgb)

            return list(font_colors)

        # Function to extract background colors from the document

        def extract_background_colors(word_file):
            document = Document(word_file)
            background_colors = set()

            for paragraph in document.paragraphs:
                for run in paragraph.runs:
                    if run.font.highlight_color is not None:
                        background_colors.add(run.font.highlight_color)

            return list(background_colors)

        # Function to extract information about bold text from the document

        def extract_bold_info(word_file):
            document = Document(word_file)
            bold_info = set()

            for paragraph in document.paragraphs:
                for run in paragraph.runs:
                    bold_info.add(run.bold)

            return list(bold_info)

        # Function to extract information about italic text from the document

        def extract_italic_info(word_file):
            document = Document(word_file)
            italic_info = set()

            for paragraph in document.paragraphs:
                for run in paragraph.runs:
                    italic_info.add(run.italic)

            return list(italic_info)

        # Function to extract information about margin right from the document

        def extract_margin_right_info(word_file):
            document = Document(word_file)
            margin_right_info = set()

            for paragraph in document.paragraphs:
                if paragraph.paragraph_format.right_indent is not None:
                    margin_right_info.add(
                        paragraph.paragraph_format.right_indent.pt)

            return list(margin_right_info)

        # Function to extract information about indent right from the document

        def extract_indent_right_info(word_file):
            document = Document(word_file)
            indent_right_info = set()

            for paragraph in document.paragraphs:
                if paragraph.paragraph_format.right_indent is not None:
                    indent_right_info.add(
                        paragraph.paragraph_format.right_indent.pt)

            return list(indent_right_info)

        # Function to extract information about underline from the document

        def extract_underline_info(word_file):
            document = Document(word_file)
            underline_info = set()

            for paragraph in document.paragraphs:
                for run in paragraph.runs:
                    underline_info.add(run.underline)

            return list(underline_info)

        # Function to extract information about page column from the document

        def extract_page_column_info(word_file):
            document = Document(word_file)
            page_column_info = set()

            for section in document.sections:
                page_column_info.add(section.start_type)

            return list(page_column_info)

        # Function to extract information about margin top from the document

        def extract_margin_top_info(word_file):
            document = Document(word_file)
            margin_top_info = set()

            for paragraph in document.paragraphs:
                if paragraph.paragraph_format.space_before is not None:
                    margin_top_info.add(
                        paragraph.paragraph_format.space_before.pt)

            return list(margin_top_info)

        # Function to extract information about margin bottom from the document
        def extract_margin_bottom_info(word_file):
            document = Document(word_file)
            margin_bottom_info = set()

            for paragraph in document.paragraphs:
                if paragraph.paragraph_format.space_after is not None:
                    margin_bottom_info.add(
                        paragraph.paragraph_format.space_after.pt)

            return list(margin_bottom_info)

        # Function to extract information about margin left from the document
        def extract_margin_left_info(word_file):
            document = Document(word_file)
            margin_left_info = set()

            for paragraph in document.paragraphs:
                if paragraph.paragraph_format.left_indent is not None:
                    margin_left_info.add(
                        paragraph.paragraph_format.left_indent.pt)

            return list(margin_left_info)

        # Function to extract information about indent left from the document
        def extract_indent_left_info(word_file):
            document = Document(word_file)
            indent_left_info = set()

            for paragraph in document.paragraphs:
                if paragraph.paragraph_format.first_line_indent is not None:
                    indent_left_info.add(
                        paragraph.paragraph_format.first_line_indent.pt)

            return list(indent_left_info)

        # Function to extract information about indent right from the document
        def extract_indent_right_info(word_file):
            document = Document(word_file)
            indent_right_info = set()

            for paragraph in document.paragraphs:
                if paragraph.paragraph_format.right_indent is not None:
                    indent_right_info.add(
                        paragraph.paragraph_format.right_indent.pt)

            return list(indent_right_info)

        # Function to extract information about spacing before from the document
        def extract_spacing_before_info(word_file):
            document = Document(word_file)
            spacing_before_info = set()

            for paragraph in document.paragraphs:
                if paragraph.paragraph_format.space_before is not None:
                    spacing_before_info.add(
                        paragraph.paragraph_format.space_before.pt)

            return list(spacing_before_info)

        # Function to extract information about spacing after from the document
        def extract_spacing_after_info(word_file):
            document = Document(word_file)
            spacing_after_info = set()

            for paragraph in document.paragraphs:
                if paragraph.paragraph_format.space_after is not None:
                    spacing_after_info.add(
                        paragraph.paragraph_format.space_after.pt)

            return list(spacing_after_info)

        # Function to extract information about alignment from the document
        def extract_alignment_info(word_file):
            document = Document(word_file)
            alignment_info = set()

            for paragraph in document.paragraphs:
                alignment_info.add(paragraph.alignment)

            return list(alignment_info)

        # Function to extract information about words like 'we', 'I', 'our', and 'you' from the document
        def extract_words_info(word_file):
            document = Document(word_file)
            words_info = set()

            for paragraph in document.paragraphs:
                for run in paragraph.runs:
                    if 'we' in run.text.lower() or 'i' in run.text.lower() or 'our' in run.text.lower() or 'you' in run.text.lower():
                        words_info.add('Yes')

            return list(words_info)

        # Function to extract information about 'Technical Point' from the document
        def extract_technical_point_info(word_file):
            document = Document(word_file)
            technical_point_info = set()

            for paragraph in document.paragraphs:
                for run in paragraph.runs:
                    if 'Technical Point' in run.text:
                        technical_point_info.add(run.text)

            return list(technical_point_info)

        # Function to extract information about 'Primary Parameters' from the document
        def extract_primary_parameters_info(word_file):
            document = Document(word_file)
            primary_parameters_info = set()

            for paragraph in document.paragraphs:
                for run in paragraph.runs:
                    if 'Primary Parameters' in run.text:
                        primary_parameters_info.add(run.text)

            return list(primary_parameters_info)

        # Function to extract information about 'Total Words' from the document
        def extract_total_words_info(word_file):
            document = Document(word_file)
            total_words_count = 0

            for paragraph in document.paragraphs:
                text = paragraph.text.strip()  
                words = text.split()            
                total_words_count += len(words) 
           
            return (total_words_count)

        # Function to extract information about 'Document Type' from the document
        def extract_document_type_info(word_file):
            kind = filetype.guess(word_file)            
            if kind is None:
                return 'it is not valid file'
            
            mime = kind.mime
            ext = kind.extension
            
            if ext == 'docx':
                document_type_info = ext       
            return document_type_info


        # Function to extract information about 'Accepted Grammar Mistake' from the document
        def extract_accepted_grammar_mistake_info(word_file):
            document = Document(word_file)
            accepted_grammar_mistake_info = set()

            for paragraph in document.paragraphs:
                for run in paragraph.runs:
                    if 'Accepted Grammar Mistake' in run.text:
                        accepted_grammar_mistake_info.add(run.text)

            return list(accepted_grammar_mistake_info)

        # Function to extract information about 'Problem discussed' from the document
        def extract_problem_discussed_info(word_file):
            document = Document(word_file)
            problem_discussed_info = set()

            for paragraph in document.paragraphs:
                for run in paragraph.runs:
                    if 'Problem discussed' in run.text:
                        problem_discussed_info.add(run.text)

            return list(problem_discussed_info)

        # Function to extract information about 'Objective' from the document
        def extract_objective_info(word_file):
            document = Document(word_file)
            objective_info = set()

            for paragraph in document.paragraphs:
                for run in paragraph.runs:
                    if 'Objective' in run.text:
                        objective_info.add(run.text)

            return list(objective_info)

        # Function to extract information about 'Novelty' from the document
        def extract_novelty_info(word_file):
            document = Document(word_file)
            novelty_info = set()

            for paragraph in document.paragraphs:
                for run in paragraph.runs:
                    if 'Novelty' in run.text:
                        novelty_info.add(run.text)

            return list(novelty_info)

        # Function to extract information about 'Technique Used' from the document
        def extract_technique_used_info(word_file):
            document = Document(word_file)
            technique_used_info = set()

            for paragraph in document.paragraphs:
                for run in paragraph.runs:
                    if 'Technique Used' in run.text:
                        technique_used_info.add(run.text)

            return list(technique_used_info)

        # Function to extract information about 'Plagiarism' from the document
        def extract_plagiarism_info(word_file):
            document = Document(word_file)
            plagiarism_info = set()

            for paragraph in document.paragraphs:
                for run in paragraph.runs:
                    if 'Plagiarism' in run.text:
                        plagiarism_info.add(run.text)

            return list(plagiarism_info)

        # Function to extract information about 'Accepted Plagiarism' from the document
        def extract_accepted_plagiarism_info(word_file):
            document = Document(word_file)
            accepted_plagiarism_info = set()

            for paragraph in document.paragraphs:
                for run in paragraph.runs:
                    if 'Accepted Plagiarism' in run.text:
                        accepted_plagiarism_info.add(run.text)

            return list(accepted_plagiarism_info)

        # Function to extract information about 'Plag Tool' from the document
        def extract_plag_tool_info(word_file):
            document = Document(word_file)
            plag_tool_info = set()

            for paragraph in document.paragraphs:
                for run in paragraph.runs:
                    if 'Plag Tool' in run.text:
                        plag_tool_info.add(run.text)

            return list(plag_tool_info)

        # Function to extract information about 'Mathematical Result' from the document

        def extract_mathematical_result_info(word_file):
            document = Document(word_file)
            mathematical_result_info = set()

            for paragraph in document.paragraphs:
                for run in paragraph.runs:
                    if 'Mathematical Result' in run.text:
                        mathematical_result_info.add(run.text)

            return list(mathematical_result_info)

        # Function to extract information about 'Result (Mathematical Description)' from the document

        def extract_mathematical_description_info(word_file):
            document = Document(word_file)
            mathematical_description_info = set()

            for paragraph in document.paragraphs:
                for run in paragraph.runs:
                    if 'Result (Mathematical Description)' in run.text:
                        mathematical_description_info.add(run.text)

            return list(mathematical_description_info)

        # Function to extract information about 'Future Scope' from the document

        def extract_future_scope_info(word_file):
            document = Document(word_file)
            future_scope_info = set()

            for paragraph in document.paragraphs:
                for run in paragraph.runs:
                    if 'Future Scope' in run.text:
                        future_scope_info.add(run.text)

            return list(future_scope_info)

        # Function to extract information about 'Dataset' from the document

        def extract_dataset_info(word_file):
            document = Document(word_file)
            dataset_info = set()

            for paragraph in document.paragraphs:
                for run in paragraph.runs:
                    if 'Dataset' in run.text:
                        dataset_info.add(run.text)

            return list(dataset_info)

        # Function to extract information about 'Dataset Description' from the document

        def extract_dataset_description_info(word_file):
            document = Document(word_file)
            dataset_description_info = set()

            for paragraph in document.paragraphs:
                for run in paragraph.runs:
                    if 'Dataset Description' in run.text:
                        dataset_description_info.add(run.text)

            return list(dataset_description_info)

        # Function to extract information about 'Dataset Source' from the document

        def extract_dataset_source_info(word_file):
            document = Document(word_file)
            dataset_source_info = set()

            for paragraph in document.paragraphs:
                for run in paragraph.runs:
                    if 'Dataset Source' in run.text:
                        dataset_source_info.add(run.text)

            return list(dataset_source_info)

        # Function to extract information about 'Dataset Type' from the document

        def extract_dataset_type_info(word_file):
            document = Document(word_file)
            dataset_type_info = set()

            for paragraph in document.paragraphs:
                for run in paragraph.runs:
                    if 'Dataset Type' in run.text:
                        dataset_type_info.add(run.text)

            return list(dataset_type_info)

        # Function to extract information about 'Keywords' from the document

        def extract_keywords_info(word_file):
            document = Document(word_file)
            keywords_info = set()

            for paragraph in document.paragraphs:
                for run in paragraph.runs:
                    if 'Keywords' in run.text:
                        keywords_info.add(run.text)

            return list(keywords_info)

        # Function to extract information about 'Keywords heading' from the document

        def extract_keywords_heading_info(word_file):
            document = Document(word_file)
            keywords_heading_info = set()

            for paragraph in document.paragraphs:
                for run in paragraph.runs:
                    if 'Keywords heading' in run.text:
                        keywords_heading_info.add(run.text)

            return list(keywords_heading_info)

        # Function to extract information about 'No. of keywords' from the document

        def extract_num_keywords_info(word_file):
            document = Document(word_file)
            num_keywords_info = set()

            for paragraph in document.paragraphs:
                for run in paragraph.runs:
                    if 'No. of keywords' in run.text:
                        num_keywords_info.add(run.text)

            return list(num_keywords_info)

        # Function to map the color index to color names

        def color_name(color_index):
            color_mapping = {
                WD_COLOR_INDEX.AUTO: "Automatic (Usually Black)",
                WD_COLOR_INDEX.BLACK: "Black",
                WD_COLOR_INDEX.BLUE: "Blue",
                WD_COLOR_INDEX.BRIGHT_GREEN: "Bright Green",
                WD_COLOR_INDEX.DARK_BLUE: "Dark Blue",
                WD_COLOR_INDEX.DARK_RED: "Dark Red",
                WD_COLOR_INDEX.DARK_YELLOW: "Dark Yellow",
                WD_COLOR_INDEX.GRAY_25: "Gray (25%)",
                WD_COLOR_INDEX.GRAY_50: "Gray (50%)",
                WD_COLOR_INDEX.GREEN: "Green",
                WD_COLOR_INDEX.PINK: "Pink",
                WD_COLOR_INDEX.RED: "Red",
                WD_COLOR_INDEX.TEAL: "Teal",
                WD_COLOR_INDEX.TURQUOISE: "Turquoise",
                WD_COLOR_INDEX.VIOLET: "Violet",
                WD_COLOR_INDEX.WHITE: "White",
                WD_COLOR_INDEX.YELLOW: "Yellow"
            }
            return color_mapping.get(color_index, "Other")

        # Function to extract background colors from the document

        def extract_highlighted_colors(word_file):
            document = Document(word_file)
            highlighted_colors = set()

            for paragraph in document.paragraphs:
                for run in paragraph.runs:
                    if run.font.highlight_color is not None:
                        color_index = run.font.highlight_color
                        color = color_name(color_index)
                        highlighted_colors.add(color)

            return highlighted_colors


        # Extract information for each specified key
        font_styles = extract_font_styles(word_file)
        font_sizes = extract_font_sizes(word_file)
        font_colors = extract_font_colors(word_file)
        background_colors = extract_background_colors(word_file)
        bold_info = extract_bold_info(word_file)
        italic_info = extract_italic_info(word_file)
        underline_info = extract_underline_info(word_file)
        page_column_info = extract_page_column_info(word_file)
        spacing_before_info = extract_spacing_before_info(word_file)
        spacing_after_info = extract_spacing_after_info(word_file)
        margin_top_info = extract_margin_top_info(word_file)
        margin_left_info = extract_margin_left_info(word_file)
        margin_right_info = extract_margin_right_info(word_file)
        margin_bottom_info = extract_margin_bottom_info(word_file)
        indent_right_info = extract_indent_right_info(word_file)
        indent_left_info = extract_indent_left_info(word_file)
        alignment_info = extract_alignment_info(word_file)
        words_info = extract_words_info(word_file)
        total_info = extract_total_words_info(word_file)
        doc_type_info = extract_document_type_info(word_file)

        # Define the output dictionary
        output_dict = {
            "Font Style": font_styles,
            "Font Size": font_sizes,
            "Font Color": font_colors,
            "Background Color": background_colors,
            "Bold": bold_info,
            "Italic": italic_info,
            "Underline": underline_info,
            "Page Column": page_column_info,
            "Spacing Before": spacing_before_info,
            "Spacing After": spacing_after_info,
            "Margin Top": margin_top_info,
            "Margin Left": margin_left_info,
            "Margin Right": margin_right_info,
            "Margin Bottom": margin_bottom_info,
            "Indent Right": indent_right_info,
            "Indent Left": indent_left_info,
            "Alignment": alignment_info,
            "Words like we, I, our and you": words_info,
            "Abstract Word Count": total_info,
            "Document Type": doc_type_info,
        }

        # Function to analyze a Word document and extract specified properties
        def analyze_word_document(word_file, output_dict):
            # Load the Word document
            doc = Document(word_file)

            # Initialize a dictionary to store property values
            properties = {}

            # Get font colors
            font_colors = {
                j.font.color.rgb for i in doc.paragraphs for j in i.runs if j.font.color and j.font.color.rgb}

            # Extract highlighted colors as background colors in the document
            background_colors = extract_highlighted_colors(word_file)

            # Define the words to check for
            words_to_check = ["we", "i", "our", "you"]

            # Check if any of the words to check for are present in the document
            found_words = set()
            for paragraph in doc.paragraphs:
                for word in words_to_check:
                    if word in paragraph.text.lower():
                        found_words.add(word)

            # Iterate through the JSON data to fetch property names and values
            for property_name in filtered_keys:
                # Initialize the set for this property
                property_set = set()

                # Handle "Indent Right" and "Indent Left" properties
                if property_name == "Indent Right" or property_name == "Indent Left":
                    value = None
                    for paragraph in doc.paragraphs:
                        if property_name == "Indent Right":
                            value = paragraph.paragraph_format.right_indent
                        elif property_name == "Indent Left":
                            value = paragraph.paragraph_format.left_indent

                        # Handle the case where the value is not None
                        if value is not None:
                            # Convert to string
                            property_set.add(str(value.pt))

                # Handle other properties
                else:
                    # Iterate through paragraphs and runs to collect unique formatting values
                    for paragraph in doc.paragraphs:
                        for run in paragraph.runs:
                            font = run.font

                            # Handle the specific property
                            if property_name == "Font Style" and font.name and font.name != "None":
                                property_set.add(font.name)
                            elif property_name == "Font Size" and font.size and font.size.pt:
                                # Convert to string
                                property_set.add(str(int(font.size.pt)))
                            elif property_name == "Bold":
                                property_set.add("Yes" if font.bold else "No")
                            elif property_name == "Italic":
                                property_set.add(
                                    "Yes" if font.italic else "No")
                            elif property_name == "Underline":
                                property_set.add(
                                    "Yes" if font.underline else "No")
                            elif property_name == "Font Color" and font.color and font.color.rgb:
                                if hasattr(font.color.rgb, 'rgb_str'):
                                    property_set.add(font.color.rgb.rgb_str())
                            elif property_name == "Background Color":
                                property_set.update(background_colors)

                    # Special handling for additional properties
                    if property_name == "Page Column":
                        property_set.add("One Column")

                    if property_name in ["Spacing Before", "Spacing After", "Margin Top", "Margin Left",
                                         "Margin Right", "Margin Bottom"]:
                        for paragraph in doc.paragraphs:
                            value = None
                            if property_name == "Spacing Before":
                                value = paragraph.paragraph_format.space_before
                            elif property_name == "Spacing After":
                                value = paragraph.paragraph_format.space_after
                            elif property_name == "Margin Top":
                                value = paragraph.paragraph_format.space_before
                            elif property_name == "Margin Left":
                                value = paragraph.paragraph_format.left_indent
                            elif property_name == "Margin Right":
                                value = paragraph.paragraph_format.right_indent
                            elif property_name == "Margin Bottom":
                                value = paragraph.paragraph_format.space_after

                            # Handle the case where the value is not None
                            if value is not None:
                                # Convert to string
                                property_set.add(str(value.pt))

                    if property_name == "Alignment":
                        for paragraph in doc.paragraphs:
                            # Convert to string
                            property_set.add(str(paragraph.alignment))

                    if property_name == "Words like we, I, our and you":
                        if found_words:
                            property_set.add("Yes")
                        else:
                            property_set.add("No")

                    if property_name == "Abstract Word Count":
                        if found_words:
                            property_set.add(total_info)
                        else:
                            property_set.add("No")

                    if property_name == "Document Type":
                        if found_words:
                            property_set.add(doc_type_info)
                        else:
                            property_set.add("No")

                    if property_name == "Font Color":
                        for color in font_colors:
                            property_set.add(color)

                # Filter out "None" values
                property_set.discard("None")

                # Convert sets to lists for JSON serialization
                properties[property_name] = list(
                    property_set) if property_set else []

            return properties

        def format_json(data, level=0):
            result = ""
            if isinstance(data, dict):
                result += "{\n"
                for key, value in data.items():
                    result += '    ' * level + f'"{key}": '
                    result += format_json(value, level + 1)
                    result += ",\n"
                result = result[:-2]  # Remove the trailing comma and newline
                result += "\n" + '    ' * (level - 1) + "}"
            elif isinstance(data, list):
                result += "["
                if data:  # Check if the list is not empty
                    for item in data:
                        result += format_json(item, level + 1) + ", "
                    result = result[:-2]  # Remove the trailing comma and space
                result += "]"
            else:
                result += f'"{data}"'
            return result

        # Call the analyze_word_document function
        result = analyze_word_document(word_file, filtered_keys)

        # Format the JSON result
        formatted_result = json.loads(format_json(result))

        # Dynamically generated comments
        comment_pool = MasterModule.objects.get(module_name=checkpoints.module_name).comment_json

        return render(request, 'dashboard/create_review.html', {'Title': 'Review', 'review_type': 'Standard', 'document_path': docx_file, 'General': general, 'Technical': technical, "auto_data": formatted_result, 'requestDetail': request_name, 'review_comment': comment_pool})
    
    return redirect('show_all_review')


# 2. Create review - Custom (formally user created, derived from Master review)
def my_custom(request):
    create_custom_review()


'''
Get the all custom requests created 
by user associated with the selected module.
'''
@login_required(login_url='userLogin')
def getCustomRequestForReview(request, id):
    request_with_module = RequestModule.objects.filter(
        username=request.user, ModuleName_id=id)
    data = serializers.serialize('json', request_with_module)
    return JsonResponse({'success': True, 'message': 'Request saved successfully', 'data': data})


# Save review explicitly
@csrf_exempt
def save_review_report(request):
    if request.method == 'POST':
        try:
            request_name = request.POST.get('request_name')
            docx_file = request.POST.get('docx_file')
            processed_file = request.POST.get('processed_file')
            json_field = request.POST.get('json_field')
            review_type = request.POST.get('reviewtype')
            review_name = request.POST.get('review_name')
            existing_review = ReviewModule.objects.filter(
                username=request.user, review_name=review_name).exists()

            if not existing_review:
                document = ReviewModule(username=request.user, request_name=request_name, docx_file=docx_file,
                                        processed_file=processed_file, json_field=json.loads(json_field), review_type=review_type, review_name=review_name)
                document.save()
                return JsonResponse({"success": True, "message": "Review saved successfully"}, status=200)
            else:
                return JsonResponse({'success': False, 'message': 'Review name already exist'})

        except Exception as e:
            return JsonResponse({"error": f"Error processing data: {str(e)}"}, status=500)
    else:
        return JsonResponse({"error": "Invalid request method"}, status=405)


# Show All Reviews
@login_required(login_url='userLogin')
def show_all_review(request):
    modulesList = MasterModule.objects.all()
    page = request.GET.get('page')
    fav = request.GET.get('fav')
    search_query = request.GET.get('search')
    date_query = request.GET.get('dateRange')

    reviewList = ReviewModule.objects.filter(
        username=request.user).order_by('-created_at')

    if search_query:
        reviewList = reviewList.filter(Q(review_name__icontains=search_query))

    if date_query:
        try:
            if date_query != None:
                start_date_str, end_date_str = date_query.split(" - ")
                formatted_start_date = timezone.datetime.strptime(
                    start_date_str, "%d/%m/%Y")
                formatted_end_date = timezone.datetime.strptime(
                    end_date_str, "%d/%m/%Y")
                reviewList = reviewList.filter(created_at__range=[
                                               formatted_start_date, formatted_end_date + timezone.timedelta(days=1)])
        except ValueError:
            pass

    if fav == 'true':
        reviewList = reviewList.filter(is_favorite=True)

    items_per_page = 3
    paginator = Paginator(reviewList, items_per_page)
    page = request.GET.get('page')

    if request.META.get('HTTP_X_REQUESTED_WITH') == 'XMLHttpRequest':
        try:
            reviewList = paginator.page(page)
        except PageNotAnInteger:
            reviewList = paginator.page(1)
        except EmptyPage:
            reviewList = paginator.page(paginator.num_pages)

        serialized_review_list = serializers.serialize('json', reviewList)
        all_pages = reviewList.paginator.num_pages
        current_page = reviewList.number

        has_previous = reviewList.has_previous()
        if has_previous:
            previous_page_number = reviewList.previous_page_number()
        else:
            previous_page_number = False

        has_next = reviewList.has_next()
        if has_next:
            next_page_number = reviewList.next_page_number()
        else:
            next_page_number = False

        data = {
            'reviewList': serialized_review_list,
            'has_previous': has_previous,
            'has_next': has_next,
            'previous_page_number': previous_page_number,
            'next_page_number': next_page_number,
            'all_pages': all_pages,
            'current_page': current_page,
        }

        return JsonResponse(data)
    return render(request, 'dashboard/show_all_review.html', {'Title': 'Review', 'modulesList': modulesList, "reviewList": reviewList})


# Delete review
@login_required(login_url='userLogin')
def delete_review(request, id):
    ReviewModule.objects.get(id=id).delete()
    return redirect('show_all_review')

# Delete all review
@csrf_exempt
@login_required
def delete_all_review(request):
    if request.method == 'POST':
        selected_ids = request.POST.getlist('selected_ids[]')
        ReviewModule.objects.filter(
            username=request.user, id__in=selected_ids).delete()

        return JsonResponse({'message': 'Reviews deleted successfully'})
    return JsonResponse({'error': 'Invalid request method'}, status=400)


# Update review
@login_required(login_url='userLogin')
def update_review(request, review_id):
    if request.method == 'POST':
        originalName = request.POST.get('originalName')
        json_field = request.POST.get('json_field')
        review_name = request.POST.get('review_name')

        obj = ReviewModule.objects.get(id=review_id)
        if review_name == originalName:
            obj.json_field = json.loads(json_field)
            obj.save()
            return JsonResponse({'success': True, 'message': 'Review Updated successfully'})
        else:
            existing_review = ReviewModule.objects.filter(username=request.user, review_name=review_name).exists()
            if existing_review == False:
                obj.review_name = review_name
                obj.json_field = json.loads(json_field)
                obj.save()
                return JsonResponse({'success': True, "message": "Data Updated successfully"}, status=200)
            else:
                return JsonResponse({'success': False, 'message': 'Review name already exist'})

    review = ReviewModule.objects.get(id=review_id)
    review_img_pool = {
        "Font": ["login-main1.png"],
    }

    return render(request, 'dashboard/update_review.html', {'Title': 'Review', 'fake_img': review_img_pool, 'review': review})


@csrf_exempt
def create_review_pdf(request):
    # id = request.POST.get(24)
    # review_instance = ReviewModule.objects.get(id=24)
    # json_data = review_instance.json_field
    # req_name = review_instance.request_name
    # rev_type = review_instance.review_type
    # context = {
    #     'data': json_data,
    #     'req_name': req_name,
    #     'rev_type': rev_type,
    # }
    if request.method == "POST":
        id = request.POST.get('id')
        review_instance = ReviewModule.objects.get(id=id)
        json_data = review_instance.json_field
        req_name = review_instance.request_name
        rev_type = review_instance.review_type
        context = {
            'data': json_data,
            'req_name': req_name,
            'rev_type': rev_type,
        }
        html_string = render_to_string('dashboard/reviewpdf.html', context)
        pdf = HTML(string=html_string).write_pdf()
        pdf_base64 = base64.b64encode(pdf).decode('utf-8')
        response_data = {
            'requestname': review_instance.review_name,
            'clientname': request.user.username,
            'pdf_data': pdf_base64,
        }

        return JsonResponse(response_data)
    return render(request, 'dashboard/reviewpdf.html', context)


@csrf_exempt
def add_favorite_request(request):
    try:
        if request.method == "POST":
            id = request.POST.get('id')
            condition = request.POST.get('isFavorite')
            favorite = condition.lower() == 'false'
            if favorite != True:
                fav_instance = RequestModule.objects.get(id=id)
                fav_instance.is_favorite = True
                fav_instance.save()
                return JsonResponse({"message": "add to Favorite", 'status': 201})
            else:
                fav_instance = RequestModule.objects.get(id=id)
                fav_instance.is_favorite = False
                fav_instance.save()
                return JsonResponse({"message": "remove from Favorite", 'status': 200})

    except Exception as e:
        return JsonResponse({'error': 'An unexpected error occurred'}, status=500)


@csrf_exempt
def add_favorite_review(request):
    try:
        if request.method == "POST":
            id = request.POST.get('id')
            condition = request.POST.get('isFavorite')
            favorite = condition.lower() == 'false'
            if favorite != True:
                fav_instance = ReviewModule.objects.get(id=id)
                fav_instance.is_favorite = True
                fav_instance.save()
                return JsonResponse({"message": "add to Favorite", 'status': 201})
            else:
                fav_instance = ReviewModule.objects.get(id=id)
                fav_instance.is_favorite = False
                fav_instance.save()
                return JsonResponse({"message": "remove from Favorite", 'status': 200})

    except Exception as e:
        return JsonResponse({'error': 'An unexpected error occurred'}, status=500)
