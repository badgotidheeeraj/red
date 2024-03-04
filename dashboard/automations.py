from .models import UserAccount, MasterModule, RequestModule, Resource
from django.contrib.auth.decorators import login_required 
from django.shortcuts import get_object_or_404, render
from docx.enum.text import WD_COLOR_INDEX, WD_PARAGRAPH_ALIGNMENT
from docx import Document
from django.shortcuts import render, redirect
import filetype
import json


# @login_required(login_url='userLogin')
# def create_master_review(request):
#     if request.method == 'POST':
#         docx_file = request.FILES['docx_file']
#         request_name = request.POST.get('request_name',) 

#         checkpoints = MasterModule.objects.get(id=request_name) 
#         json_data = checkpoints.json_field
#         general = json_data['General Point']
#         technical = json_data['Technical Point']

#         exclude_keys = ["General Point", "Technical Point"] 
#         keys = []
         
#         def extract_keys_recursive(dictionary): 
#             for key, value in dictionary.items():
#                 keys.append(key)  # Include the key itself 
#                 if isinstance(value, dict):
#                     extract_keys_recursive(value)  # Recurse into nested dictionary
#         extract_keys_recursive(json_data)
        
#         exclude_keys = ["General Point", "Font", "Page Layout"]
#         filtered_keys = [key for key in keys if key not in exclude_keys]
#         word_file = docx_file
#         # Define the keys and filtered_keys based on your requirements
#         filtered_keys = ['Font Style', 'Font Size', 'Font Color', 'Background Color', 'Bold', 'Italic', 'Underline', 'Page Column', 'Spacing Before', 'Spacing After', 'Margin Top', 'Margin Left', 'Margin Right', 'Margin Bottom', 'Indent Right', 'Indent Left', 'Alignment', 'Words like we, I, our and you']
        
#         # Function to extract font styles from the document
#         def extract_font_styles(word_file):
#             document = Document(word_file)
#             font_styles = set()

#             for paragraph in document.paragraphs:
#                 for run in paragraph.runs:
#                     font_styles.add(run.font.name)

#             return list(font_styles)

#         # Function to extract font sizes from the document
#         def extract_font_sizes(word_file):
#             document = Document(word_file)
#             font_sizes = set()

#             for paragraph in document.paragraphs:
#                 for run in paragraph.runs:
#                     if run.font.size is not None:
#                         font_sizes.add(run.font.size.pt)

#             return list(font_sizes)


#         # Function to extract font colors from the document
#         def extract_font_colors(word_file):
#             document = Document(word_file)
#             font_colors = set()

#             for paragraph in document.paragraphs:
#                 for run in paragraph.runs:
#                     font_colors.add(run.font.color.rgb)

#             return list(font_colors)
        

#         # Function to extract background colors from the document
#         def extract_background_colors(word_file):
#             document = Document(word_file)
#             background_colors = set()

#             for paragraph in document.paragraphs:
#                 for run in paragraph.runs:
#                     if run.font.highlight_color is not None:
#                         background_colors.add(run.font.highlight_color)

#             return list(background_colors)


#         # Function to extract information about bold text from the document
#         def extract_bold_info(word_file):
#             document = Document(word_file)
#             bold_info = set()

#             for paragraph in document.paragraphs:
#                 for run in paragraph.runs:
#                     bold_info.add(run.bold)

#             return list(bold_info)

        
#         # Function to extract information about italic text from the document
#         def extract_italic_info(word_file):
#             document = Document(word_file)
#             italic_info = set()

#             for paragraph in document.paragraphs:
#                 for run in paragraph.runs:
#                     italic_info.add(run.italic)

#             return list(italic_info)


#         # Function to extract information about margin right from the document
#         def extract_margin_right_info(word_file):
#             document = Document(word_file)
#             margin_right_info = set()

#             for paragraph in document.paragraphs:
#                 if paragraph.paragraph_format.right_indent is not None:
#                     margin_right_info.add(paragraph.paragraph_format.right_indent.pt)

#             return list(margin_right_info)


#         # Function to extract information about indent right from the document
#         def extract_indent_right_info(word_file):
#             document = Document(word_file)
#             indent_right_info = set()

#             for paragraph in document.paragraphs:
#                 if paragraph.paragraph_format.right_indent is not None:
#                     indent_right_info.add(paragraph.paragraph_format.right_indent.pt)

#             return list(indent_right_info)

        
#         # Function to extract information about underline from the document
#         def extract_underline_info(word_file):
#             document = Document(word_file)
#             underline_info = set()

#             for paragraph in document.paragraphs:
#                 for run in paragraph.runs:
#                     underline_info.add(run.underline)

#             return list(underline_info)


#         # Function to extract information about page column from the document
#         def extract_page_column_info(word_file):
#             document = Document(word_file)
#             page_column_info = set()

#             for section in document.sections:
#                 page_column_info.add(section.start_type)

#             return list(page_column_info)


#         # Function to extract information about margin top from the document
#         def extract_margin_top_info(word_file):
#             document = Document(word_file)
#             margin_top_info = set()

#             for paragraph in document.paragraphs:
#                 if paragraph.paragraph_format.space_before is not None:
#                     margin_top_info.add(paragraph.paragraph_format.space_before.pt)

#             return list(margin_top_info)

#         # Function to extract information about margin bottom from the document
#         def extract_margin_bottom_info(word_file):
#             document = Document(word_file)
#             margin_bottom_info = set()

#             for paragraph in document.paragraphs:
#                 if paragraph.paragraph_format.space_after is not None:
#                     margin_bottom_info.add(paragraph.paragraph_format.space_after.pt)

#             return list(margin_bottom_info)

#         # Function to extract information about margin left from the document
#         def extract_margin_left_info(word_file):
#             document = Document(word_file)
#             margin_left_info = set()

#             for paragraph in document.paragraphs:
#                 if paragraph.paragraph_format.left_indent is not None:
#                     margin_left_info.add(paragraph.paragraph_format.left_indent.pt)

#             return list(margin_left_info)

#         # Function to extract information about indent left from the document
#         def extract_indent_left_info(word_file):
#             document = Document(word_file)
#             indent_left_info = set()

#             for paragraph in document.paragraphs:
#                 if paragraph.paragraph_format.first_line_indent is not None:
#                     indent_left_info.add(paragraph.paragraph_format.first_line_indent.pt)

#             return list(indent_left_info)

#         # Function to extract information about indent right from the document
#         def extract_indent_right_info(word_file):
#             document = Document(word_file)
#             indent_right_info = set()

#             for paragraph in document.paragraphs:
#                 if paragraph.paragraph_format.right_indent is not None:
#                     indent_right_info.add(paragraph.paragraph_format.right_indent.pt)

#             return list(indent_right_info)

#         # Function to extract information about spacing before from the document
#         def extract_spacing_before_info(word_file):
#             document = Document(word_file)
#             spacing_before_info = set()

#             for paragraph in document.paragraphs:
#                 if paragraph.paragraph_format.space_before is not None:
#                     spacing_before_info.add(paragraph.paragraph_format.space_before.pt)

#             return list(spacing_before_info)

#         # Function to extract information about spacing after from the document
#         def extract_spacing_after_info(word_file):
#             document = Document(word_file)
#             spacing_after_info = set()

#             for paragraph in document.paragraphs:
#                 if paragraph.paragraph_format.space_after is not None:
#                     spacing_after_info.add(paragraph.paragraph_format.space_after.pt)

#             return list(spacing_after_info)

#         # Function to extract information about alignment from the document
#         def extract_alignment_info(word_file):
#             document = Document(word_file)
#             alignment_info = set()

#             for paragraph in document.paragraphs:
#                 alignment_info.add(paragraph.alignment)

#             return list(alignment_info)


#         # Function to extract information about words like 'we', 'I', 'our', and 'you' from the document
#         def extract_words_info(word_file):
#             document = Document(word_file)
#             words_info = set()

#             for paragraph in document.paragraphs:
#                 for run in paragraph.runs:
#                     if 'we' in run.text.lower() or 'i' in run.text.lower() or 'our' in run.text.lower() or 'you' in run.text.lower():
#                         words_info.add('Yes')

#             return list(words_info)

        
#         # Function to extract information about 'Technical Point' from the document
#         def extract_technical_point_info(word_file):
#             document = Document(word_file)
#             technical_point_info = set()

#             for paragraph in document.paragraphs:
#                 for run in paragraph.runs:
#                     if 'Technical Point' in run.text:
#                         technical_point_info.add(run.text)

#             return list(technical_point_info)


#         # Function to extract information about 'Primary Parameters' from the document
#         def extract_primary_parameters_info(word_file):
#             document = Document(word_file)
#             primary_parameters_info = set()

#             for paragraph in document.paragraphs:
#                 for run in paragraph.runs:
#                     if 'Primary Parameters' in run.text:
#                         primary_parameters_info.add(run.text)

#             return list(primary_parameters_info)


#         # Function to extract information about 'Total Words' from the document
#         def extract_total_words_info(word_file):
#             document = Document(word_file)
#             total_words_info = set()

#             for paragraph in document.paragraphs:
#                 for run in paragraph.runs:
#                     if 'Total Words' in run.text:
#                         total_words_info.add(run.text)

#             return list(total_words_info)


#         # Function to extract information about 'Document Type' from the document
#         def extract_document_type_info(word_file):
#             document = Document(word_file)
#             document_type_info = set()

#             for paragraph in document.paragraphs:
#                 for run in paragraph.runs:
#                     if 'Document Type' in run.text:
#                         document_type_info.add(run.text)

#             return list(document_type_info)


#         # Function to extract information about 'Accepted Grammar Mistake' from the document
#         def extract_accepted_grammar_mistake_info(word_file):
#             document = Document(word_file)
#             accepted_grammar_mistake_info = set()

#             for paragraph in document.paragraphs:
#                 for run in paragraph.runs:
#                     if 'Accepted Grammar Mistake' in run.text:
#                         accepted_grammar_mistake_info.add(run.text)

#             return list(accepted_grammar_mistake_info)


#         # Function to extract information about 'Problem discussed' from the document
#         def extract_problem_discussed_info(word_file):
#             document = Document(word_file)
#             problem_discussed_info = set()

#             for paragraph in document.paragraphs:
#                 for run in paragraph.runs:
#                     if 'Problem discussed' in run.text:
#                         problem_discussed_info.add(run.text)

#             return list(problem_discussed_info)


#         # Function to extract information about 'Objective' from the document
#         def extract_objective_info(word_file):
#             document = Document(word_file)
#             objective_info = set()

#             for paragraph in document.paragraphs:
#                 for run in paragraph.runs:
#                     if 'Objective' in run.text:
#                         objective_info.add(run.text)

#             return list(objective_info)


#         # Function to extract information about 'Novelty' from the document
#         def extract_novelty_info(word_file):
#             document = Document(word_file)
#             novelty_info = set()

#             for paragraph in document.paragraphs:
#                 for run in paragraph.runs:
#                     if 'Novelty' in run.text:
#                         novelty_info.add(run.text)

#             return list(novelty_info)


#         # Function to extract information about 'Technique Used' from the document
#         def extract_technique_used_info(word_file):
#             document = Document(word_file)
#             technique_used_info = set()

#             for paragraph in document.paragraphs:
#                 for run in paragraph.runs:
#                     if 'Technique Used' in run.text:
#                         technique_used_info.add(run.text)

#             return list(technique_used_info)


#         # Function to extract information about 'Plagiarism' from the document
#         def extract_plagiarism_info(word_file):
#             document = Document(word_file)
#             plagiarism_info = set()

#             for paragraph in document.paragraphs:
#                 for run in paragraph.runs:
#                     if 'Plagiarism' in run.text:
#                         plagiarism_info.add(run.text)

#             return list(plagiarism_info)


#         # Function to extract information about 'Accepted Plagiarism' from the document
#         def extract_accepted_plagiarism_info(word_file):
#             document = Document(word_file)
#             accepted_plagiarism_info = set()

#             for paragraph in document.paragraphs:
#                 for run in paragraph.runs:
#                     if 'Accepted Plagiarism' in run.text:
#                         accepted_plagiarism_info.add(run.text)

#             return list(accepted_plagiarism_info)


#         # Function to extract information about 'Plag Tool' from the document
#         def extract_plag_tool_info(word_file):
#             document = Document(word_file)
#             plag_tool_info = set()

#             for paragraph in document.paragraphs:
#                 for run in paragraph.runs:
#                     if 'Plag Tool' in run.text:
#                         plag_tool_info.add(run.text)

#             return list(plag_tool_info)


#         # Function to extract information about 'Mathematical Result' from the document
#         def extract_mathematical_result_info(word_file):
#             document = Document(word_file)
#             mathematical_result_info = set()

#             for paragraph in document.paragraphs:
#                 for run in paragraph.runs:
#                     if 'Mathematical Result' in run.text:
#                         mathematical_result_info.add(run.text)

#             return list(mathematical_result_info)


#         # Function to extract information about 'Result (Mathematical Description)' from the document
#         def extract_mathematical_description_info(word_file):
#             document = Document(word_file)
#             mathematical_description_info = set()

#             for paragraph in document.paragraphs:
#                 for run in paragraph.runs:
#                     if 'Result (Mathematical Description)' in run.text:
#                         mathematical_description_info.add(run.text)

#             return list(mathematical_description_info)

        
#         # Function to extract information about 'Future Scope' from the document
#         def extract_future_scope_info(word_file):
#             document = Document(word_file)
#             future_scope_info = set()

#             for paragraph in document.paragraphs:
#                 for run in paragraph.runs:
#                     if 'Future Scope' in run.text:
#                         future_scope_info.add(run.text)

#             return list(future_scope_info)


#         # Function to extract information about 'Dataset' from the document
#         def extract_dataset_info(word_file):
#             document = Document(word_file)
#             dataset_info = set()

#             for paragraph in document.paragraphs:
#                 for run in paragraph.runs:
#                     if 'Dataset' in run.text:
#                         dataset_info.add(run.text)

#             return list(dataset_info)


#         # Function to extract information about 'Dataset Description' from the document
#         def extract_dataset_description_info(word_file):
#             document = Document(word_file)
#             dataset_description_info = set()

#             for paragraph in document.paragraphs:
#                 for run in paragraph.runs:
#                     if 'Dataset Description' in run.text:
#                         dataset_description_info.add(run.text)

#             return list(dataset_description_info)


#         # Function to extract information about 'Dataset Source' from the document
#         def extract_dataset_source_info(word_file):
#             document = Document(word_file)
#             dataset_source_info = set()

#             for paragraph in document.paragraphs:
#                 for run in paragraph.runs:
#                     if 'Dataset Source' in run.text:
#                         dataset_source_info.add(run.text)

#             return list(dataset_source_info)


#         # Function to extract information about 'Dataset Type' from the document
#         def extract_dataset_type_info(word_file):
#             document = Document(word_file)
#             dataset_type_info = set()

#             for paragraph in document.paragraphs:
#                 for run in paragraph.runs:
#                     if 'Dataset Type' in run.text:
#                         dataset_type_info.add(run.text)

#             return list(dataset_type_info)


#         # Function to extract information about 'Keywords' from the document
#         def extract_keywords_info(word_file):
#             document = Document(word_file)
#             keywords_info = set()

#             for paragraph in document.paragraphs:
#                 for run in paragraph.runs:
#                     if 'Keywords' in run.text:
#                         keywords_info.add(run.text)

#             return list(keywords_info)


#         # Function to extract information about 'Keywords heading' from the document
#         def extract_keywords_heading_info(word_file):
#             document = Document(word_file)
#             keywords_heading_info = set()

#             for paragraph in document.paragraphs:
#                 for run in paragraph.runs:
#                     if 'Keywords heading' in run.text:
#                         keywords_heading_info.add(run.text)

#             return list(keywords_heading_info)


#         # Function to extract information about 'No. of keywords' from the document
#         def extract_num_keywords_info(word_file):
#             document = Document(word_file)
#             num_keywords_info = set()

#             for paragraph in document.paragraphs:
#                 for run in paragraph.runs:
#                     if 'No. of keywords' in run.text:
#                         num_keywords_info.add(run.text)

#             return list(num_keywords_info)

        
#         # Function to map the color index to color names
#         def color_name(color_index):
#             color_mapping = {
#                 WD_COLOR_INDEX.AUTO: "Automatic (Usually Black)",
#                 WD_COLOR_INDEX.BLACK: "Black",
#                 WD_COLOR_INDEX.BLUE: "Blue",
#                 WD_COLOR_INDEX.BRIGHT_GREEN: "Bright Green",
#                 WD_COLOR_INDEX.DARK_BLUE: "Dark Blue",
#                 WD_COLOR_INDEX.DARK_RED: "Dark Red",
#                 WD_COLOR_INDEX.DARK_YELLOW: "Dark Yellow",
#                 WD_COLOR_INDEX.GRAY_25: "Gray (25%)",
#                 WD_COLOR_INDEX.GRAY_50: "Gray (50%)",
#                 WD_COLOR_INDEX.GREEN: "Green",
#                 WD_COLOR_INDEX.PINK: "Pink",
#                 WD_COLOR_INDEX.RED: "Red",
#                 WD_COLOR_INDEX.TEAL: "Teal",
#                 WD_COLOR_INDEX.TURQUOISE: "Turquoise",
#                 WD_COLOR_INDEX.VIOLET: "Violet",
#                 WD_COLOR_INDEX.WHITE: "White",
#                 WD_COLOR_INDEX.YELLOW: "Yellow"
#             }
#             return color_mapping.get(color_index, "Other")

        
#         # Function to extract background colors from the document
#         def extract_highlighted_colors(word_file):
#             document = Document(word_file)
#             highlighted_colors = set()

#             for paragraph in document.paragraphs:
#                 for run in paragraph.runs:
#                     if run.font.highlight_color is not None:
#                         color_index = run.font.highlight_color
#                         color = color_name(color_index)
#                         highlighted_colors.add(color)

#             return highlighted_colors

        
#         # Recursive function to extract keys from a nested dictionary
#         def extract_keys_recursive(dictionary):
#             keys = []
#             for key, value in dictionary.items():
#                 keys.append(key)  # Include the key itself
#                 if isinstance(value, dict):
#                     keys.extend(extract_keys_recursive(value))  # Recurse into nested dictionary
#             return keys

#         word_file = docx_file

#         # Extract information for each specified key
#         font_styles = extract_font_styles(word_file)
#         font_sizes = extract_font_sizes(word_file)
#         font_colors = extract_font_colors(word_file)
#         background_colors = extract_background_colors(word_file)
#         bold_info = extract_bold_info(word_file)
#         italic_info = extract_italic_info(word_file)
#         underline_info = extract_underline_info(word_file)
#         page_column_info = extract_page_column_info(word_file)
#         spacing_before_info = extract_spacing_before_info(word_file)
#         spacing_after_info = extract_spacing_after_info(word_file)
#         margin_top_info = extract_margin_top_info(word_file)
#         margin_left_info = extract_margin_left_info(word_file)
#         margin_right_info = extract_margin_right_info(word_file)
#         margin_bottom_info = extract_margin_bottom_info(word_file)
#         indent_right_info = extract_indent_right_info(word_file)
#         indent_left_info = extract_indent_left_info(word_file)
#         alignment_info = extract_alignment_info(word_file)
#         words_info = extract_words_info(word_file)

#         # Define the output dictionary
#         output_dict = {
#             "Font Style": font_styles,
#             "Font Size": font_sizes,
#             "Font Color": font_colors,
#             "Background Color": background_colors,
#             "Bold": bold_info,
#             "Italic": italic_info,
#             "Underline": underline_info,
#             "Page Column": page_column_info,
#             "Spacing Before": spacing_before_info,
#             "Spacing After": spacing_after_info,
#             "Margin Top": margin_top_info,
#             "Margin Left": margin_left_info,
#             "Margin Right": margin_right_info,
#             "Margin Bottom": margin_bottom_info,
#             "Indent Right": indent_right_info,
#             "Indent Left": indent_left_info,
#             "Alignment": alignment_info,
#             "Words like we, I, our and you": words_info
#         }

#         # Function to analyze a Word document and extract specified properties
#         def analyze_word_document(word_file, output_dict):
#             # Load the Word document
#             doc = Document(word_file)

#             # Initialize a dictionary to store property values
#             properties = {}

#             # Get font colors
#             font_colors = {j.font.color.rgb for i in doc.paragraphs for j in i.runs if j.font.color and j.font.color.rgb}

#             # Extract highlighted colors as background colors in the document
#             background_colors = extract_highlighted_colors(word_file)

#             # Define the words to check for
#             words_to_check = ["we", "i", "our", "you"]
            
#             # Check if any of the words to check for are present in the document
#             found_words = set()
#             for paragraph in doc.paragraphs:
#                 for word in words_to_check:
#                     if word in paragraph.text.lower():
#                         found_words.add(word)

#             # Iterate through the JSON data to fetch property names and values
#             for property_name in filtered_keys:
#                 # Initialize the set for this property
#                 property_set = set()

#                 # Handle "Indent Right" and "Indent Left" properties
#                 if property_name == "Indent Right" or property_name == "Indent Left":
#                     value = None
#                     for paragraph in doc.paragraphs:
#                         if property_name == "Indent Right":
#                             value = paragraph.paragraph_format.right_indent
#                         elif property_name == "Indent Left":
#                             value = paragraph.paragraph_format.left_indent

#                         # Handle the case where the value is not None
#                         if value is not None:
#                             property_set.add(str(value.pt))  # Convert to string

#                 # Handle other properties
#                 else:
#             # Iterate through paragraphs and runs to collect unique formatting values
#                     for paragraph in doc.paragraphs:
#                         for run in paragraph.runs:
#                             font = run.font

#                             # Handle the specific property
#                             if property_name == "Font Style" and font.name and font.name != "None":
#                                 property_set.add(font.name)
#                             elif property_name == "Font Size" and font.size and font.size.pt:
#                                 property_set.add(str(int(font.size.pt)))  # Convert to string
#                             elif property_name == "Bold":
#                                 property_set.add("Yes" if font.bold else "No")
#                             elif property_name == "Italic":
#                                 property_set.add("Yes" if font.italic else "No")
#                             elif property_name == "Underline":
#                                 property_set.add("Yes" if font.underline else "No")
#                             elif property_name == "Font Color" and font.color and font.color.rgb:
#                                 if hasattr(font.color.rgb, 'rgb_str'):
#                                     property_set.add(font.color.rgb.rgb_str())
#                             elif property_name == "Background Color":
#                                 property_set.update(background_colors)


#                     # Special handling for additional properties
#                     if property_name == "Page Column":
#                         property_set.add("One Column")

#                     if property_name in ["Spacing Before", "Spacing After", "Margin Top", "Margin Left",
#                                 "Margin Right", "Margin Bottom"]:
#                         for paragraph in doc.paragraphs:
#                             value = None
#                             if property_name == "Spacing Before":
#                                 value = paragraph.paragraph_format.space_before
#                             elif property_name == "Spacing After":
#                                 value = paragraph.paragraph_format.space_after
#                             elif property_name == "Margin Top":
#                                 value = paragraph.paragraph_format.space_before
#                             elif property_name == "Margin Left":
#                                 value = paragraph.paragraph_format.left_indent
#                             elif property_name == "Margin Right":
#                                 value = paragraph.paragraph_format.right_indent
#                             elif property_name == "Margin Bottom":
#                                 value = paragraph.paragraph_format.space_after

#                             # Handle the case where the value is not None
#                             if value is not None:
#                                 property_set.add(str(value.pt))  # Convert to string

#                     if property_name == "Alignment":
#                         for paragraph in doc.paragraphs:
#                             property_set.add(str(paragraph.alignment))  # Convert to string

#                     if property_name == "Words like we, I, our and you":
#                         if found_words:
#                             property_set.add("Yes")
#                         else:
#                             property_set.add("No")
#                     if property_name == "Font Color":
#                         for color in font_colors:
#                             property_set.add(color)

#                 # Filter out "None" values
#                 property_set.discard("None")

#                 # Convert sets to lists for JSON serialization
#                 properties[property_name] = list(property_set) if property_set else []

#             return properties

        
#         def format_json(data, level=0):
#             result = ""
#             if isinstance(data, dict):
#                 result += "{\n"
#                 for key, value in data.items():
#                     result += '    ' * level + f'"{key}": '
#                     result += format_json(value, level + 1)
#                     result += ",\n"
#                 result = result[:-2]  # Remove the trailing comma and newline
#                 result += "\n" + '    ' * (level - 1) + "}"
#             elif isinstance(data, list):
#                 result += "["
#                 if data:  # Check if the list is not empty
#                     for item in data:
#                         result += format_json(item, level + 1) + ", "
#                     result = result[:-2]  # Remove the trailing comma and space
#                 result += "]"
#             else:
#                 result += f'"{data}"'
#             return result
 
#         # Call the analyze_word_document function
#         result = analyze_word_document(word_file, filtered_keys)

#         # Format the JSON result 
#         formatted_result = json.loads(format_json(result))


#         # Dynamically generated comments
#         comment_pool = MasterModule.objects.get(module_name=checkpoints.module_name).comment_json

#         return render(request, 'dashboard/create_review.html', {'Title':'Review','review_type': 'Standard','document_path': docx_file, 'General': general, 'Technical': technical, "auto_data": formatted_result, 'requestDetail': request_name, 'review_comment':comment_pool,})


@login_required(login_url='userLogin')
def create_custom_review(request):
    if request.method == 'POST':
        docx_file = request.FILES['docx_file'] 
        request_name = request.POST.get('request_name',) 
        req = RequestModule.objects.get(id=request_name) 
        json_data = req.json_field  
        general = json_data['General Point']
        technical = json_data['Technical Point']

        exclude_keys = ["General Point", "Technical Point"] 
        keys = []
         
        def extract_keys_recursive(dictionary): 
            for key, value in dictionary.items():
                keys.append(key)  # Include the key itself 
                if isinstance(value, dict):
                    extract_keys_recursive(value)  # Recurse into nested dictionary 
        extract_keys_recursive(general)

        def extract_keys_recursive(dictionary): 
            for key, value in dictionary.items():
                keys.append(key)  # Include the key itself 
                if isinstance(value, dict):
                    extract_keys_recursive(value)  # Recurse into nested dictionary
        extract_keys_recursive(technical)
        
        exclude_keys = ["General Point", "Font", "Page Layout","Technical Point","Additional Parameters","Abstract should have","Dataset","Keywords"]
        filtered_keys = [key for key in keys if key not in exclude_keys] 
        word_file = docx_file  
        # Define the keys and filtered_keys based on your requirements
        filtered_keys = ['Font Style', 'Font Size', 'Font Color', 'Background Color', 'Bold', 'Italic', 'Underline', 'Page Column', 'Spacing Before', 'Spacing After', 'Margin Top', 'Margin Left', 'Margin Right', 'Margin Bottom', 'Indent Right', 'Indent Left', 'Alignment', 'Words like we, I, our and you','Abstract Word Count','Document Type','Accepted Grammar Mistakes (Syntactically)','Accepted Plagiarism','Plag Tool','Problem Discussed','Objective','Novelty','Technique Used','Result (Mathematical Description)','Future Scope','Dataset Description','Dataset Source','Dataset Type','Keywords Heading','No. of keywords']
        
        # Function to extract font styles from the document
        def extract_font_styles(word_file):
            document = Document(word_file)
            font_styles = set()

            for paragraph in document.paragraphs:
                for run in paragraph.runs:
                    font_styles.add(run.font.name)

            return list(font_styles)


        # Function to extract font sizes from the document
        def extract_font_sizes(word_file):
            document = Document(word_file)
            font_sizes = set()

            for paragraph in document.paragraphs:
                for run in paragraph.runs:
                    if run.font.size is not None:
                        font_sizes.add(run.font.size.pt)

            return list(font_sizes)


        # Function to extract font colors from the document
        def extract_font_colors(word_file):
            document = Document(word_file)
            font_colors = set()

            for paragraph in document.paragraphs:
                for run in paragraph.runs:
                    font_colors.add(run.font.color.rgb)

            return list(font_colors)
        

        # Function to extract background colors from the document
        def extract_background_colors(word_file):
            document = Document(word_file)
            background_colors = set()

            for paragraph in document.paragraphs:
                for run in paragraph.runs:
                    if run.font.highlight_color is not None:
                        background_colors.add(run.font.highlight_color)

            return list(background_colors)


        # Function to extract information about bold text from the document
        def extract_bold_info(word_file):
            document = Document(word_file)
            bold_info = set()

            for paragraph in document.paragraphs:
                for run in paragraph.runs:
                    bold_info.add(run.bold)

            return list(bold_info)

        
        # Function to extract information about italic text from the document
        def extract_italic_info(word_file):
            document = Document(word_file)
            italic_info = set()

            for paragraph in document.paragraphs:
                for run in paragraph.runs:
                    italic_info.add(run.italic)

            return list(italic_info)


        # Function to extract information about margin right from the document
        def extract_margin_right_info(word_file):
            document = Document(word_file)
            margin_right_info = set()

            for paragraph in document.paragraphs:
                if paragraph.paragraph_format.right_indent is not None:
                    margin_right_info.add(paragraph.paragraph_format.right_indent.pt)

            return list(margin_right_info)


        # Function to extract information about indent right from the document
        def extract_indent_right_info(word_file):
            document = Document(word_file)
            indent_right_info = set()

            for paragraph in document.paragraphs:
                if paragraph.paragraph_format.right_indent is not None:
                    indent_right_info.add(paragraph.paragraph_format.right_indent.pt)

            return list(indent_right_info)

        
        # Function to extract information about underline from the document
        def extract_underline_info(word_file):
            document = Document(word_file)
            underline_info = set()

            for paragraph in document.paragraphs:
                for run in paragraph.runs:
                    underline_info.add(run.underline)

            return list(underline_info)


        # Function to extract information about page column from the document
        def extract_page_column_info(word_file):
            document = Document(word_file)
            page_column_info = set()

            for section in document.sections:
                page_column_info.add(section.start_type)

            return list(page_column_info)


        # Function to extract information about margin top from the document
        def extract_margin_top_info(word_file):
            document = Document(word_file)
            margin_top_info = set()

            for paragraph in document.paragraphs:
                if paragraph.paragraph_format.space_before is not None:
                    margin_top_info.add(paragraph.paragraph_format.space_before.pt)

            return list(margin_top_info)

        # Function to extract information about margin bottom from the document
        def extract_margin_bottom_info(word_file):
            document = Document(word_file)
            margin_bottom_info = set()

            for paragraph in document.paragraphs:
                if paragraph.paragraph_format.space_after is not None:
                    margin_bottom_info.add(paragraph.paragraph_format.space_after.pt)

            return list(margin_bottom_info)

        # Function to extract information about margin left from the document
        def extract_margin_left_info(word_file):
            document = Document(word_file)
            margin_left_info = set()

            for paragraph in document.paragraphs:
                if paragraph.paragraph_format.left_indent is not None:
                    margin_left_info.add(paragraph.paragraph_format.left_indent.pt)

            return list(margin_left_info)

        # Function to extract information about indent left from the document
        def extract_indent_left_info(word_file):
            document = Document(word_file)
            indent_left_info = set()

            for paragraph in document.paragraphs:
                if paragraph.paragraph_format.first_line_indent is not None:
                    indent_left_info.add(paragraph.paragraph_format.first_line_indent.pt)

            return list(indent_left_info)

        # Function to extract information about indent right from the document
        def extract_indent_right_info(word_file):
            document = Document(word_file)
            indent_right_info = set()

            for paragraph in document.paragraphs:
                if paragraph.paragraph_format.right_indent is not None:
                    indent_right_info.add(paragraph.paragraph_format.right_indent.pt)

            return list(indent_right_info)

        # Function to extract information about spacing before from the document
        def extract_spacing_before_info(word_file):
            document = Document(word_file)
            spacing_before_info = set()

            for paragraph in document.paragraphs:
                if paragraph.paragraph_format.space_before is not None:
                    spacing_before_info.add(paragraph.paragraph_format.space_before.pt)

            return list(spacing_before_info)

        # Function to extract information about spacing after from the document
        def extract_spacing_after_info(word_file):
            document = Document(word_file)
            spacing_after_info = set()

            for paragraph in document.paragraphs:
                if paragraph.paragraph_format.space_after is not None:
                    spacing_after_info.add(paragraph.paragraph_format.space_after.pt)

            return list(spacing_after_info)

        # Function to extract information about alignment from the document
        def extract_alignment_info(word_file):
            document = Document(word_file)
            alignment_info = set()

            for paragraph in document.paragraphs:
                alignment_info.add(paragraph.alignment)

            return list(alignment_info)


        # Function to extract information about words like 'we', 'I', 'our', and 'you' from the document
        def extract_words_info(word_file):
            document = Document(word_file)
            words_info = set()

            for paragraph in document.paragraphs:
                for run in paragraph.runs:
                    if 'we' in run.text.lower() or 'i' in run.text.lower() or 'our' in run.text.lower() or 'you' in run.text.lower():
                        words_info.add('Yes')

            return list(words_info)

        
        # Function to extract information about 'Technical Point' from the document
        def extract_technical_point_info(word_file):
            document = Document(word_file)
            technical_point_info = set()

            for paragraph in document.paragraphs:
                for run in paragraph.runs:
                    if 'Technical Point' in run.text:
                        technical_point_info.add(run.text)

            return list(technical_point_info)


        # Function to extract information about 'Primary Parameters' from the document
        def extract_primary_parameters_info(word_file):
            document = Document(word_file)
            primary_parameters_info = set()

            for paragraph in document.paragraphs:
                for run in paragraph.runs:
                    if 'Primary Parameters' in run.text:
                        primary_parameters_info.add(run.text)

            return list(primary_parameters_info)


        # Function to extract information about 'Total Words' from the document
        def extract_total_words_info(word_file):
            document = Document(word_file)
            total_words_count = 0

            for paragraph in document.paragraphs:
                text = paragraph.text.strip()  
                words = text.split()            
                total_words_count += len(words) 
           
            return (total_words_count)



        # Function to extract information about 'Document Type' from the document
        def extract_document_type_info(word_file):
            kind = filetype.guess(word_file)            
            if kind is None:
                return 'it is not valid file'
            
            mime = kind.mime
            ext = kind.extension
            
            if ext == 'docx':
                document_type_info = ext       
            return document_type_info


        # Function to extract information about 'Accepted Grammar Mistake' from the document
        def extract_accepted_grammar_mistake_info(word_file):
            document = Document(word_file)
            accepted_grammar_mistake_info = set()

            for paragraph in document.paragraphs:
                for run in paragraph.runs:
                    if 'Accepted Grammar Mistake' in run.text:
                        accepted_grammar_mistake_info.add(run.text)

            return list(accepted_grammar_mistake_info)


        # Function to extract information about 'Problem discussed' from the document
        def extract_problem_discussed_info(word_file):
            document = Document(word_file)
            problem_discussed_info = set()

            for paragraph in document.paragraphs:
                for run in paragraph.runs:
                    if 'Problem discussed' in run.text:
                        problem_discussed_info.add(run.text)

            return list(problem_discussed_info)


        # Function to extract information about 'Objective' from the document
        def extract_objective_info(word_file):
            document = Document(word_file)
            objective_info = set()

            for paragraph in document.paragraphs:
                for run in paragraph.runs:
                    if 'Objective' in run.text:
                        objective_info.add(run.text)

            return list(objective_info)


        # Function to extract information about 'Novelty' from the document
        def extract_novelty_info(word_file):
            document = Document(word_file)
            novelty_info = set()

            for paragraph in document.paragraphs:
                for run in paragraph.runs:
                    if 'Novelty' in run.text:
                        novelty_info.add(run.text)

            return list(novelty_info)


        # Function to extract information about 'Technique Used' from the document
        def extract_technique_used_info(word_file):
            document = Document(word_file)
            technique_used_info = set()

            for paragraph in document.paragraphs:
                for run in paragraph.runs:
                    if 'Technique Used' in run.text:
                        technique_used_info.add(run.text)

            return list(technique_used_info)


        # Function to extract information about 'Plagiarism' from the document
        def extract_plagiarism_info(word_file):
            document = Document(word_file)
            plagiarism_info = set()

            for paragraph in document.paragraphs:
                for run in paragraph.runs:
                    if 'Plagiarism' in run.text:
                        plagiarism_info.add(run.text)

            return list(plagiarism_info)


        # Function to extract information about 'Accepted Plagiarism' from the document
        def extract_accepted_plagiarism_info(word_file):
            document = Document(word_file)
            accepted_plagiarism_info = set()

            for paragraph in document.paragraphs:
                for run in paragraph.runs:
                    if 'Accepted Plagiarism' in run.text:
                        accepted_plagiarism_info.add(run.text)

            return list(accepted_plagiarism_info)


        # Function to extract information about 'Plag Tool' from the document
        def extract_plag_tool_info(word_file):
            document = Document(word_file)
            plag_tool_info = set()

            for paragraph in document.paragraphs:
                for run in paragraph.runs:
                    if 'Plag Tool' in run.text:
                        plag_tool_info.add(run.text)

            return list(plag_tool_info)


        # Function to extract information about 'Mathematical Result' from the document
        def extract_mathematical_result_info(word_file):
            document = Document(word_file)
            mathematical_result_info = set()

            for paragraph in document.paragraphs:
                for run in paragraph.runs:
                    if 'Mathematical Result' in run.text:
                        mathematical_result_info.add(run.text)

            return list(mathematical_result_info)


        # Function to extract information about 'Result (Mathematical Description)' from the document
        def extract_mathematical_description_info(word_file):
            document = Document(word_file)
            mathematical_description_info = set()

            for paragraph in document.paragraphs:
                for run in paragraph.runs:
                    if 'Result (Mathematical Description)' in run.text:
                        mathematical_description_info.add(run.text)

            return list(mathematical_description_info)

        
        # Function to extract information about 'Future Scope' from the document
        def extract_future_scope_info(word_file):
            document = Document(word_file)
            future_scope_info = set()

            for paragraph in document.paragraphs:
                for run in paragraph.runs:
                    if 'Future Scope' in run.text:
                        future_scope_info.add(run.text)

            return list(future_scope_info)


        # Function to extract information about 'Dataset' from the document
        def extract_dataset_info(word_file):
            document = Document(word_file)
            dataset_info = set()

            for paragraph in document.paragraphs:
                for run in paragraph.runs:
                    if 'Dataset' in run.text:
                        dataset_info.add(run.text)

            return list(dataset_info)


        # Function to extract information about 'Dataset Description' from the document
        def extract_dataset_description_info(word_file):
            document = Document(word_file)
            dataset_description_info = set()

            for paragraph in document.paragraphs:
                for run in paragraph.runs:
                    if 'Dataset Description' in run.text:
                        dataset_description_info.add(run.text)

            return list(dataset_description_info)


        # Function to extract information about 'Dataset Source' from the document
        def extract_dataset_source_info(word_file):
            document = Document(word_file)
            dataset_source_info = set()

            for paragraph in document.paragraphs:
                for run in paragraph.runs:
                    if 'Dataset Source' in run.text:
                        dataset_source_info.add(run.text)

            return list(dataset_source_info)


        # Function to extract information about 'Dataset Type' from the document
        def extract_dataset_type_info(word_file):
            document = Document(word_file)
            dataset_type_info = set()

            for paragraph in document.paragraphs:
                for run in paragraph.runs:
                    if 'Dataset Type' in run.text:
                        dataset_type_info.add(run.text)

            return list(dataset_type_info)


        # Function to extract information about 'Keywords' from the document
        def extract_keywords_info(word_file):
            document = Document(word_file)
            keywords_info = set()

            for paragraph in document.paragraphs:
                for run in paragraph.runs:
                    if 'Keywords' in run.text:
                        keywords_info.add(run.text)

            return list(keywords_info)


        # Function to extract information about 'Keywords heading' from the document
        def extract_keywords_heading_info(word_file):
            document = Document(word_file)
            keywords_heading_info = set()

            for paragraph in document.paragraphs:
                for run in paragraph.runs:
                    if 'Keywords heading' in run.text:
                        keywords_heading_info.add(run.text)

            return list(keywords_heading_info)


        # Function to extract information about 'No. of keywords' from the document
        def extract_num_keywords_info(word_file):
            document = Document(word_file)
            num_keywords_info = set()

            for paragraph in document.paragraphs:
                for run in paragraph.runs:
                    if 'No. of keywords' in run.text:
                        num_keywords_info.add(run.text)

            return list(num_keywords_info)

        
        # Function to map the color index to color names
        def color_name(color_index):
            color_mapping = {
                WD_COLOR_INDEX.AUTO: "Automatic (Usually Black)",
                WD_COLOR_INDEX.BLACK: "Black",
                WD_COLOR_INDEX.BLUE: "Blue",
                WD_COLOR_INDEX.BRIGHT_GREEN: "Bright Green",
                WD_COLOR_INDEX.DARK_BLUE: "Dark Blue",
                WD_COLOR_INDEX.DARK_RED: "Dark Red",
                WD_COLOR_INDEX.DARK_YELLOW: "Dark Yellow",
                WD_COLOR_INDEX.GRAY_25: "Gray (25%)",
                WD_COLOR_INDEX.GRAY_50: "Gray (50%)",
                WD_COLOR_INDEX.GREEN: "Green",
                WD_COLOR_INDEX.PINK: "Pink",
                WD_COLOR_INDEX.RED: "Red",
                WD_COLOR_INDEX.TEAL: "Teal",
                WD_COLOR_INDEX.TURQUOISE: "Turquoise",
                WD_COLOR_INDEX.VIOLET: "Violet",
                WD_COLOR_INDEX.WHITE: "White",
                WD_COLOR_INDEX.YELLOW: "Yellow"
            }
            return color_mapping.get(color_index, "Other")

        
        # Function to extract background colors from the document
        def extract_highlighted_colors(word_file):
            document = Document(word_file)
            highlighted_colors = set()

            for paragraph in document.paragraphs:
                for run in paragraph.runs:
                    if run.font.highlight_color is not None:
                        color_index = run.font.highlight_color
                        color = color_name(color_index)
                        highlighted_colors.add(color)

            return highlighted_colors

        
        # Recursive function to extract keys from a nested dictionary
        def extract_keys_recursive(dictionary):
            keys = []
            for key, value in dictionary.items():
                keys.append(key)  # Include the key itself
                if isinstance(value, dict):
                    keys.extend(extract_keys_recursive(value))  # Recurse into nested dictionary
            return keys

        word_file = docx_file

        # Extract information for each specified key
        font_styles = extract_font_styles(word_file)
        font_sizes = extract_font_sizes(word_file)
        font_colors = extract_font_colors(word_file)
        background_colors = extract_background_colors(word_file)
        bold_info = extract_bold_info(word_file)
        italic_info = extract_italic_info(word_file)
        underline_info = extract_underline_info(word_file)
        page_column_info = extract_page_column_info(word_file)
        spacing_before_info = extract_spacing_before_info(word_file)
        spacing_after_info = extract_spacing_after_info(word_file)
        margin_top_info = extract_margin_top_info(word_file)
        margin_left_info = extract_margin_left_info(word_file)
        margin_right_info = extract_margin_right_info(word_file)
        margin_bottom_info = extract_margin_bottom_info(word_file)
        indent_right_info = extract_indent_right_info(word_file)
        indent_left_info = extract_indent_left_info(word_file)
        alignment_info = extract_alignment_info(word_file)
        words_info = extract_words_info(word_file)
        total_info = extract_total_words_info(word_file)
        doc_type_info = extract_document_type_info(word_file)

        # Define the output dictionary
        output_dict = {
            "Font Style": font_styles,
            "Font Size": font_sizes,
            "Font Color": font_colors,
            "Background Color": background_colors,
            "Bold": bold_info,
            "Italic": italic_info,
            "Underline": underline_info,
            "Page Column": page_column_info,
            "Spacing Before": spacing_before_info,
            "Spacing After": spacing_after_info,
            "Margin Top": margin_top_info,
            "Margin Left": margin_left_info,
            "Margin Right": margin_right_info,
            "Margin Bottom": margin_bottom_info,
            "Indent Right": indent_right_info,
            "Indent Left": indent_left_info,
            "Alignment": alignment_info,
            "Words like we, I, our and you": words_info,
            "Abstract Word Count": total_info,
            "Document Type": doc_type_info,
        }

        # Function to analyze a Word document and extract specified properties
        def analyze_word_document(word_file, output_dict):
            # Load the Word document
            doc = Document(word_file)

            # Initialize a dictionary to store property values
            properties = {}

            # Get font colors
            font_colors = {j.font.color.rgb for i in doc.paragraphs for j in i.runs if j.font.color and j.font.color.rgb}

            # Extract highlighted colors as background colors in the document
            background_colors = extract_highlighted_colors(word_file)

            # Define the words to check for
            words_to_check = ["we", "i", "our", "you"]
            
            # Check if any of the words to check for are present in the document
            found_words = set()
            for paragraph in doc.paragraphs:
                for word in words_to_check:
                    if word in paragraph.text.lower():
                        found_words.add(word)

            # Iterate through the JSON data to fetch property names and values
            for property_name in filtered_keys:
                # Initialize the set for this property
                property_set = set()

                # Handle "Indent Right" and "Indent Left" properties
                if property_name == "Indent Right" or property_name == "Indent Left":
                    value = None
                    for paragraph in doc.paragraphs:
                        if property_name == "Indent Right":
                            value = paragraph.paragraph_format.right_indent
                        elif property_name == "Indent Left":
                            value = paragraph.paragraph_format.left_indent

                        # Handle the case where the value is not None
                        if value is not None:
                            property_set.add(str(value.pt))  # Convert to string

                # Handle other properties
                else:
            # Iterate through paragraphs and runs to collect unique formatting values
                    for paragraph in doc.paragraphs:
                        for run in paragraph.runs:
                            font = run.font

                            # Handle the specific property
                            if property_name == "Font Style" and font.name and font.name != "None":
                                property_set.add(font.name)
                            elif property_name == "Font Size" and font.size and font.size.pt:
                                property_set.add(str(int(font.size.pt)))  # Convert to string
                            elif property_name == "Bold":
                                property_set.add("Yes" if font.bold else "No")
                            elif property_name == "Italic":
                                property_set.add("Yes" if font.italic else "No")
                            elif property_name == "Underline":
                                property_set.add("Yes" if font.underline else "No")
                            elif property_name == "Font Color" and font.color and font.color.rgb:
                                if hasattr(font.color.rgb, 'rgb_str'):
                                    property_set.add(font.color.rgb.rgb_str())
                            elif property_name == "Background Color":
                                property_set.update(background_colors)


                    # Special handling for additional properties
                    if property_name == "Page Column":
                        property_set.add("One Column")

                    if property_name in ["Spacing Before", "Spacing After", "Margin Top", "Margin Left",
                                "Margin Right", "Margin Bottom"]:
                        for paragraph in doc.paragraphs:
                            value = None
                            if property_name == "Spacing Before":
                                value = paragraph.paragraph_format.space_before
                            elif property_name == "Spacing After":
                                value = paragraph.paragraph_format.space_after
                            elif property_name == "Margin Top":
                                value = paragraph.paragraph_format.space_before
                            elif property_name == "Margin Left":
                                value = paragraph.paragraph_format.left_indent
                            elif property_name == "Margin Right":
                                value = paragraph.paragraph_format.right_indent
                            elif property_name == "Margin Bottom":
                                value = paragraph.paragraph_format.space_after

                            # Handle the case where the value is not None
                            if value is not None:
                                property_set.add(str(value.pt))  # Convert to string

                    if property_name == "Alignment":
                        for paragraph in doc.paragraphs:
                            property_set.add(str(paragraph.alignment))  # Convert to string

                    if property_name == "Words like we, I, our and you":
                        if found_words:
                            property_set.add("Yes")
                        else:
                            property_set.add("No")

                    if property_name == "Abstract Word Count":
                        if found_words:
                            property_set.add(total_info)
                        else:
                            property_set.add("No")

                    if property_name == "Document Type":
                        if found_words:
                            property_set.add(doc_type_info)
                        else:
                            property_set.add("No")
                            
                    if property_name == "Font Color":
                        for color in font_colors:
                            property_set.add(color)

                # Filter out "None" values
                property_set.discard("None")

                # Convert sets to lists for JSON serialization
                properties[property_name] = list(property_set) if property_set else []

            return properties

        
        def format_json(data, level=0):
            result = ""
            if isinstance(data, dict):
                result += "{\n"
                for key, value in data.items():
                    result += '    ' * level + f'"{key}": '
                    result += format_json(value, level + 1)
                    result += ",\n"
                result = result[:-2]  # Remove the trailing comma and newline
                result += "\n" + '    ' * (level - 1) + "}"
            elif isinstance(data, list):
                result += "["
                if data:  # Check if the list is not empty
                    for item in data:
                        result += format_json(item, level + 1) + ", "
                    result = result[:-2]  # Remove the trailing comma and space
                result += "]"
            else:
                result += f'"{data}"'
            return result
 
        # Call the analyze_word_document function
        result = analyze_word_document(word_file, filtered_keys)

        # Format the JSON result
        formatted_result = json.loads(format_json(result))

        # Dynamically generated comments
        comment_pool = RequestModule.objects.get(request_name=req.request_name).ModuleName.comment_json

        return render(request, 'dashboard/create_review.html', {'Title':'Review','review_type': 'Custom','document_path': docx_file, 'General': general, 'Technical': technical, "auto_data": formatted_result, 'requestDetail': request_name,"req_name":req,'review_comment':comment_pool})
    
    return redirect('show_all_review')



 






















































































































# ==================================================== Don't touch here, it is original code ==================================================================


# Function to extract font styles from the document
# def extract_font_styles(word_file):
#     document = Document(word_file)
#     font_styles = set()

#     for paragraph in document.paragraphs:
#         for run in paragraph.runs:
#             font_styles.add(run.font.name)

#     return list(font_styles)

# Function to extract font sizes from the document
# def extract_font_sizes(word_file):
#     document = Document(word_file)
#     font_sizes = set()

#     for paragraph in document.paragraphs:
#         for run in paragraph.runs:
#             if run.font.size is not None:
#                 font_sizes.add(run.font.size.pt)

#     return list(font_sizes)


# Function to extract font colors from the document
# def extract_font_colors(word_file):
#     document = Document(word_file)
#     font_colors = set()

#     for paragraph in document.paragraphs:
#         for run in paragraph.runs:
#             font_colors.add(run.font.color.rgb)

#     return list(font_colors)


# Function to extract background colors from the document
# def extract_background_colors(word_file):
#     document = Document(word_file)
#     background_colors = set()

#     for paragraph in document.paragraphs:
#         for run in paragraph.runs:
#             if run.font.highlight_color is not None:
#                 background_colors.add(run.font.highlight_color)

#     return list(background_colors)


# Function to extract information about bold text from the document
# def extract_bold_info(word_file):
#     document = Document(word_file)
#     bold_info = set()

#     for paragraph in document.paragraphs:
#         for run in paragraph.runs:
#             bold_info.add(run.bold)

#     return list(bold_info)


# Function to extract information about italic text from the document
# def extract_italic_info(word_file):
#     document = Document(word_file)
#     italic_info = set()

#     for paragraph in document.paragraphs:
#         for run in paragraph.runs:
#             italic_info.add(run.italic)

#     return list(italic_info)


# Function to extract information about margin right from the document
# def extract_margin_right_info(word_file):
#     document = Document(word_file)
#     margin_right_info = set()

#     for paragraph in document.paragraphs:
#         if paragraph.paragraph_format.right_indent is not None:
#             margin_right_info.add(paragraph.paragraph_format.right_indent.pt)

#     return list(margin_right_info)


# Function to extract information about indent right from the document
# def extract_indent_right_info(word_file):
#     document = Document(word_file)
#     indent_right_info = set()

#     for paragraph in document.paragraphs:
#         if paragraph.paragraph_format.right_indent is not None:
#             indent_right_info.add(paragraph.paragraph_format.right_indent.pt)

#     return list(indent_right_info)


# Function to extract information about underline from the document
# def extract_underline_info(word_file):
#     document = Document(word_file)
#     underline_info = set()

#     for paragraph in document.paragraphs:
#         for run in paragraph.runs:
#             underline_info.add(run.underline)

#     return list(underline_info)


# Function to extract information about page column from the document
# def extract_page_column_info(word_file):
#     document = Document(word_file)
#     page_column_info = set()

#     for section in document.sections:
#         page_column_info.add(section.start_type)

#     return list(page_column_info)


# Function to extract information about margin top from the document
# def extract_margin_top_info(word_file):
#     document = Document(word_file)
#     margin_top_info = set()

#     for paragraph in document.paragraphs:
#         if paragraph.paragraph_format.space_before is not None:
#             margin_top_info.add(paragraph.paragraph_format.space_before.pt)

#     return list(margin_top_info)

# Function to extract information about margin bottom from the document
# def extract_margin_bottom_info(word_file):
#     document = Document(word_file)
#     margin_bottom_info = set()

#     for paragraph in document.paragraphs:
#         if paragraph.paragraph_format.space_after is not None:
#             margin_bottom_info.add(paragraph.paragraph_format.space_after.pt)

#     return list(margin_bottom_info)

# Function to extract information about margin left from the document
# def extract_margin_left_info(word_file):
#     document = Document(word_file)
#     margin_left_info = set()

#     for paragraph in document.paragraphs:
#         if paragraph.paragraph_format.left_indent is not None:
#             margin_left_info.add(paragraph.paragraph_format.left_indent.pt)

#     return list(margin_left_info)

# Function to extract information about indent left from the document
# def extract_indent_left_info(word_file):
#     document = Document(word_file)
#     indent_left_info = set()

#     for paragraph in document.paragraphs:
#         if paragraph.paragraph_format.first_line_indent is not None:
#             indent_left_info.add(paragraph.paragraph_format.first_line_indent.pt)

#     return list(indent_left_info)

# Function to extract information about indent right from the document
# def extract_indent_right_info(word_file):
#     document = Document(word_file)
#     indent_right_info = set()

#     for paragraph in document.paragraphs:
#         if paragraph.paragraph_format.right_indent is not None:
#             indent_right_info.add(paragraph.paragraph_format.right_indent.pt)

#     return list(indent_right_info)

# Function to extract information about spacing before from the document
# def extract_spacing_before_info(word_file):
#     document = Document(word_file)
#     spacing_before_info = set()

#     for paragraph in document.paragraphs:
#         if paragraph.paragraph_format.space_before is not None:
#             spacing_before_info.add(paragraph.paragraph_format.space_before.pt)

#     return list(spacing_before_info)

# Function to extract information about spacing after from the document
# def extract_spacing_after_info(word_file):
#     document = Document(word_file)
#     spacing_after_info = set()

#     for paragraph in document.paragraphs:
#         if paragraph.paragraph_format.space_after is not None:
#             spacing_after_info.add(paragraph.paragraph_format.space_after.pt)

#     return list(spacing_after_info)

# Function to extract information about alignment from the document
# def extract_alignment_info(word_file):
#     document = Document(word_file)
#     alignment_info = set()

#     for paragraph in document.paragraphs:
#         alignment_info.add(paragraph.alignment)

#     return list(alignment_info)


# Function to extract information about words like 'we', 'I', 'our', and 'you' from the document
# def extract_words_info(word_file):
#     document = Document(word_file)
#     words_info = set()

#     for paragraph in document.paragraphs:
#         for run in paragraph.runs:
#             if 'we' in run.text.lower() or 'i' in run.text.lower() or 'our' in run.text.lower() or 'you' in run.text.lower():
#                 words_info.add('Yes')

#     return list(words_info)


# Function to extract information about 'Technical Point' from the document
# def extract_technical_point_info(word_file):
#     document = Document(word_file)
#     technical_point_info = set()

#     for paragraph in document.paragraphs:
#         for run in paragraph.runs:
#             if 'Technical Point' in run.text:
#                 technical_point_info.add(run.text)

#     return list(technical_point_info)


# Function to extract information about 'Primary Parameters' from the document
# def extract_primary_parameters_info(word_file):
#     document = Document(word_file)
#     primary_parameters_info = set()

#     for paragraph in document.paragraphs:
#         for run in paragraph.runs:
#             if 'Primary Parameters' in run.text:
#                 primary_parameters_info.add(run.text)

#     return list(primary_parameters_info)


# # Function to extract information about 'Total Words' from the document
# # def extract_total_words_info(word_file):
#     document = Document(word_file)
#     total_words_info = set()

#     for paragraph in document.paragraphs:
#         for run in paragraph.runs:
#             if 'Total Words' in run.text:
#                 total_words_info.add(run.text)

#     return list(total_words_info)


# Function to extract information about 'Document Type' from the document
# def extract_document_type_info(word_file):
#     document = Document(word_file)
#     document_type_info = set()

#     for paragraph in document.paragraphs:
#         for run in paragraph.runs:
#             if 'Document Type' in run.text:
#                 document_type_info.add(run.text)

#     return list(document_type_info)


# Function to extract information about 'Accepted Grammar Mistake' from the document
# def extract_accepted_grammar_mistake_info(word_file):
#     document = Document(word_file)
#     accepted_grammar_mistake_info = set()

#     for paragraph in document.paragraphs:
#         for run in paragraph.runs:
#             if 'Accepted Grammar Mistake' in run.text:
#                 accepted_grammar_mistake_info.add(run.text)

#     return list(accepted_grammar_mistake_info)


# Function to extract information about 'Problem discussed' from the document
# def extract_problem_discussed_info(word_file):
#     document = Document(word_file)
#     problem_discussed_info = set()

#     for paragraph in document.paragraphs:
#         for run in paragraph.runs:
#             if 'Problem discussed' in run.text:
#                 problem_discussed_info.add(run.text)

#     return list(problem_discussed_info)


# Function to extract information about 'Objective' from the document
# def extract_objective_info(word_file):
#     document = Document(word_file)
#     objective_info = set()

#     for paragraph in document.paragraphs:
#         for run in paragraph.runs:
#             if 'Objective' in run.text:
#                 objective_info.add(run.text)

#     return list(objective_info)


# Function to extract information about 'Novelty' from the document
# def extract_novelty_info(word_file):
#     document = Document(word_file)
#     novelty_info = set()

#     for paragraph in document.paragraphs:
#         for run in paragraph.runs:
#             if 'Novelty' in run.text:
#                 novelty_info.add(run.text)

#     return list(novelty_info)


# Function to extract information about 'Technique Used' from the document
# def extract_technique_used_info(word_file):
#     document = Document(word_file)
#     technique_used_info = set()

#     for paragraph in document.paragraphs:
#         for run in paragraph.runs:
#             if 'Technique Used' in run.text:
#                 technique_used_info.add(run.text)

#     return list(technique_used_info)


# Function to extract information about 'Plagiarism' from the document
# def extract_plagiarism_info(word_file):
#     document = Document(word_file)
#     plagiarism_info = set()

#     for paragraph in document.paragraphs:
#         for run in paragraph.runs:
#             if 'Plagiarism' in run.text:
#                 plagiarism_info.add(run.text)

#     return list(plagiarism_info)


# Function to extract information about 'Accepted Plagiarism' from the document
# def extract_accepted_plagiarism_info(word_file):
#     document = Document(word_file)
#     accepted_plagiarism_info = set()

#     for paragraph in document.paragraphs:
#         for run in paragraph.runs:
#             if 'Accepted Plagiarism' in run.text:
#                 accepted_plagiarism_info.add(run.text)

#     return list(accepted_plagiarism_info)


# Function to extract information about 'Plag Tool' from the document
# def extract_plag_tool_info(word_file):
#     document = Document(word_file)
#     plag_tool_info = set()

#     for paragraph in document.paragraphs:
#         for run in paragraph.runs:
#             if 'Plag Tool' in run.text:
#                 plag_tool_info.add(run.text)

#     return list(plag_tool_info)


# Function to extract information about 'Mathematical Result' from the document
# def extract_mathematical_result_info(word_file):
#     document = Document(word_file)
#     mathematical_result_info = set()

#     for paragraph in document.paragraphs:
#         for run in paragraph.runs:
#             if 'Mathematical Result' in run.text:
#                 mathematical_result_info.add(run.text)

#     return list(mathematical_result_info)


# Function to extract information about 'Result (Mathematical Description)' from the document
# def extract_mathematical_description_info(word_file):
#     document = Document(word_file)
#     mathematical_description_info = set()

#     for paragraph in document.paragraphs:
#         for run in paragraph.runs:
#             if 'Result (Mathematical Description)' in run.text:
#                 mathematical_description_info.add(run.text)

#     return list(mathematical_description_info)


# Function to extract information about 'Future Scope' from the document
# def extract_future_scope_info(word_file):
#     document = Document(word_file)
#     future_scope_info = set()

#     for paragraph in document.paragraphs:
#         for run in paragraph.runs:
#             if 'Future Scope' in run.text:
#                 future_scope_info.add(run.text)

#     return list(future_scope_info)


# Function to extract information about 'Dataset' from the document
# def extract_dataset_info(word_file):
#     document = Document(word_file)
#     dataset_info = set()

#     for paragraph in document.paragraphs:
#         for run in paragraph.runs:
#             if 'Dataset' in run.text:
#                 dataset_info.add(run.text)

#     return list(dataset_info)


# Function to extract information about 'Dataset Description' from the document
# def extract_dataset_description_info(word_file):
#     document = Document(word_file)
#     dataset_description_info = set()

#     for paragraph in document.paragraphs:
#         for run in paragraph.runs:
#             if 'Dataset Description' in run.text:
#                 dataset_description_info.add(run.text)

#     return list(dataset_description_info)


# Function to extract information about 'Dataset Source' from the document
# def extract_dataset_source_info(word_file):
#     document = Document(word_file)
#     dataset_source_info = set()

#     for paragraph in document.paragraphs:
#         for run in paragraph.runs:
#             if 'Dataset Source' in run.text:
#                 dataset_source_info.add(run.text)

#     return list(dataset_source_info)


# Function to extract information about 'Dataset Type' from the document
# def extract_dataset_type_info(word_file):
#     document = Document(word_file)
#     dataset_type_info = set()

#     for paragraph in document.paragraphs:
#         for run in paragraph.runs:
#             if 'Dataset Type' in run.text:
#                 dataset_type_info.add(run.text)

#     return list(dataset_type_info)


# Function to extract information about 'Keywords' from the document
# def extract_keywords_info(word_file):
#     document = Document(word_file)
#     keywords_info = set()

#     for paragraph in document.paragraphs:
#         for run in paragraph.runs:
#             if 'Keywords' in run.text:
#                 keywords_info.add(run.text)

#     return list(keywords_info)


# Function to extract information about 'Keywords heading' from the document
# def extract_keywords_heading_info(word_file):
#     document = Document(word_file)
#     keywords_heading_info = set()

#     for paragraph in document.paragraphs:
#         for run in paragraph.runs:
#             if 'Keywords heading' in run.text:
#                 keywords_heading_info.add(run.text)

#     return list(keywords_heading_info)


# Function to extract information about 'No. of keywords' from the document
# def extract_num_keywords_info(word_file):
#     document = Document(word_file)
#     num_keywords_info = set()

#     for paragraph in document.paragraphs:
#         for run in paragraph.runs:
#             if 'No. of keywords' in run.text:
#                 num_keywords_info.add(run.text)

#     return list(num_keywords_info)


# Function to map the color index to color names
# def color_name(color_index):
#     color_mapping = {
#         WD_COLOR_INDEX.AUTO: "Automatic (Usually Black)",
#         WD_COLOR_INDEX.BLACK: "Black",
#         WD_COLOR_INDEX.BLUE: "Blue",
#         WD_COLOR_INDEX.BRIGHT_GREEN: "Bright Green",
#         WD_COLOR_INDEX.DARK_BLUE: "Dark Blue",
#         WD_COLOR_INDEX.DARK_RED: "Dark Red",
#         WD_COLOR_INDEX.DARK_YELLOW: "Dark Yellow",
#         WD_COLOR_INDEX.GRAY_25: "Gray (25%)",
#         WD_COLOR_INDEX.GRAY_50: "Gray (50%)",
#         WD_COLOR_INDEX.GREEN: "Green",
#         WD_COLOR_INDEX.PINK: "Pink",
#         WD_COLOR_INDEX.RED: "Red",
#         WD_COLOR_INDEX.TEAL: "Teal",
#         WD_COLOR_INDEX.TURQUOISE: "Turquoise",
#         WD_COLOR_INDEX.VIOLET: "Violet",
#         WD_COLOR_INDEX.WHITE: "White",
#         WD_COLOR_INDEX.YELLOW: "Yellow"
#     }
#     return color_mapping.get(color_index, "Other")


# Function to extract background colors from the document
# def extract_highlighted_colors(word_file):
#     document = Document(word_file)
#     highlighted_colors = set()

#     for paragraph in document.paragraphs:
#         for run in paragraph.runs:
#             if run.font.highlight_color is not None:
#                 color_index = run.font.highlight_color
#                 color = color_name(color_index)
#                 highlighted_colors.add(color)

#     return highlighted_colors


# def file_automation(word_file, request_name):
#     checkpoints = MasterModule.objects.get(id=request_name)
#     json_data = checkpoints.json_field
#     general = json_data['General Point']
#     technical = json_data['Technical Point']

#     exclude_keys = ["General Point", "Technical Point"]
#     keys = []
        
#     def extract_keys_recursive(dictionary): 
#         for key, value in dictionary.items():
#             keys.append(key)  # Include the key itself 
#             if isinstance(value, dict):
#                 extract_keys_recursive(value)  # Recurse into nested dictionary
#     extract_keys_recursive(json_data)
    
#     exclude_keys = ["General Point", "Font", "Page Layout"]
#     filtered_keys = [key for key in keys if key not in exclude_keys]

#     # Define the keys and filtered_keys based on your requirements
#     filtered_keys = ['Font Style', 'Font Size', 'Font Color', 'Background Color', 'Bold', 'Italic', 'Underline', 'Page Column', 'Spacing Before', 'Spacing After', 'Margin Top', 'Margin Left', 'Margin Right', 'Margin Bottom', 'Indent Right', 'Indent Left', 'Alignment', 'Words like we, I, our and you']
    
#     # Recursive function to extract keys from a nested dictionary
#     def extract_keys_recursive(dictionary):
#         keys = []
#         for key, value in dictionary.items():
#             keys.append(key)  # Include the key itself
#             if isinstance(value, dict):
#                 keys.extend(extract_keys_recursive(value))  # Recurse into nested dictionary
#         return keys

#     # Define the output dictionary
#     output_dict = {
#         "Font Style": extract_font_styles(word_file),
#         "Font Size": extract_font_sizes(word_file),
#         "Font Color": extract_font_colors(word_file),
#         "Background Color": extract_background_colors(word_file),
#         "Bold": extract_bold_info(word_file),
#         "Italic": extract_italic_info(word_file),
#         "Underline": extract_underline_info(word_file),
#         "Page Column": extract_page_column_info(word_file),
#         "Spacing Before": extract_spacing_before_info(word_file),
#         "Spacing After": extract_spacing_after_info(word_file),
#         "Margin Top": extract_margin_top_info(word_file),
#         "Margin Left": extract_margin_left_info(word_file),
#         "Margin Right": extract_margin_right_info(word_file),
#         "Margin Bottom": extract_margin_bottom_info(word_file),
#         "Indent Right": extract_indent_right_info(word_file),
#         "Indent Left": extract_indent_left_info(word_file),
#         "Alignment": extract_alignment_info(word_file),
#         "Words like we, I, our and you": extract_words_info(word_file)
#     }

#     # Function to analyze a Word document and extract specified properties
#     def analyze_word_document(word_file, output_dict):
#         # Load the Word document
#         doc = Document(word_file)

#         # Initialize a dictionary to store property values
#         properties = {}

#         # Get font colors
#         font_colors = {j.font.color.rgb for i in doc.paragraphs for j in i.runs if j.font.color and j.font.color.rgb}

#         # Extract highlighted colors as background colors in the document
#         background_colors = extract_highlighted_colors(word_file)

#         # Define the words to check for
#         words_to_check = ["we", "i", "our", "you"]
        
#         # Check if any of the words to check for are present in the document
#         found_words = set()
#         for paragraph in doc.paragraphs:
#             for word in words_to_check:
#                 if word in paragraph.text.lower():
#                     found_words.add(word)

#         # Iterate through the JSON data to fetch property names and values
#         for property_name in filtered_keys:
#             # Initialize the set for this property
#             property_set = set()

#             # Handle "Indent Right" and "Indent Left" properties
#             if property_name == "Indent Right" or property_name == "Indent Left":
#                 value = None
#                 for paragraph in doc.paragraphs:
#                     if property_name == "Indent Right":
#                         value = paragraph.paragraph_format.right_indent
#                     elif property_name == "Indent Left":
#                         value = paragraph.paragraph_format.left_indent

#                     # Handle the case where the value is not None
#                     if value is not None:
#                         property_set.add(str(value.pt))  # Convert to string

#             # Handle other properties
#             else:
#         # Iterate through paragraphs and runs to collect unique formatting values
#                 for paragraph in doc.paragraphs:
#                     for run in paragraph.runs:
#                         font = run.font

#                         # Handle the specific property
#                         if property_name == "Font Style" and font.name and font.name != "None":
#                             property_set.add(font.name)
#                         elif property_name == "Font Size" and font.size and font.size.pt:
#                             property_set.add(str(int(font.size.pt)))  # Convert to string
#                         elif property_name == "Bold":
#                             property_set.add("Yes" if font.bold else "No")
#                         elif property_name == "Italic":
#                             property_set.add("Yes" if font.italic else "No")
#                         elif property_name == "Underline":
#                             property_set.add("Yes" if font.underline else "No")
#                         elif property_name == "Font Color" and font.color and font.color.rgb:
#                             if hasattr(font.color.rgb, 'rgb_str'):
#                                 property_set.add(font.color.rgb.rgb_str())
#                         elif property_name == "Background Color":
#                             property_set.update(background_colors)


#                 # Special handling for additional properties
#                 if property_name == "Page Column":
#                     property_set.add("One Column")

#                 if property_name in ["Spacing Before", "Spacing After", "Margin Top", "Margin Left",
#                             "Margin Right", "Margin Bottom"]:
#                     for paragraph in doc.paragraphs:
#                         value = None
#                         if property_name == "Spacing Before":
#                             value = paragraph.paragraph_format.space_before
#                         elif property_name == "Spacing After":
#                             value = paragraph.paragraph_format.space_after
#                         elif property_name == "Margin Top":
#                             value = paragraph.paragraph_format.space_before
#                         elif property_name == "Margin Left":
#                             value = paragraph.paragraph_format.left_indent
#                         elif property_name == "Margin Right":
#                             value = paragraph.paragraph_format.right_indent
#                         elif property_name == "Margin Bottom":
#                             value = paragraph.paragraph_format.space_after

#                         # Handle the case where the value is not None
#                         if value is not None:
#                             property_set.add(str(value.pt))  # Convert to string

#                 if property_name == "Alignment":
#                     for paragraph in doc.paragraphs:
#                         property_set.add(str(paragraph.alignment))  # Convert to string

#                 if property_name == "Words like we, I, our and you":
#                     if found_words:
#                         property_set.add("Yes")
#                     else:
#                         property_set.add("No")
#                 if property_name == "Font Color":
#                     for color in font_colors:
#                         property_set.add(color)

#             # Filter out "None" values
#             property_set.discard("None")

#             # Convert sets to lists for JSON serialization
#             properties[property_name] = list(property_set) if property_set else []

#         return properties

    
#     def format_json(data, level=0):
#         result = ""
#         if isinstance(data, dict):
#             result += "{\n"
#             for key, value in data.items():
#                 result += '    ' * level + f'"{key}": '
#                 result += format_json(value, level + 1)
#                 result += ",\n"
#             result = result[:-2]  # Remove the trailing comma and newline
#             result += "\n" + '    ' * (level - 1) + "}"
#         elif isinstance(data, list):
#             result += "["
#             if data:  # Check if the list is not empty
#                 for item in data:
#                     result += format_json(item, level + 1) + ", "
#                 result = result[:-2]  # Remove the trailing comma and space
#             result += "]"
#         else:
#             result += f'"{data}"'
#         return result

#     # Call the analyze_word_document function
#     result = analyze_word_document(word_file, filtered_keys)

#     # Format the JSON result 
#     formatted_result = json.loads(format_json(result))


#     # Dynamically generated comments
#     comment_pool = MasterModule.objects.get(module_name=checkpoints.module_name).comment_json

#     context = {'General': general, 'Technical': technical, "auto_data": formatted_result, 'requestDetail': request_name, 'review_comment':comment_pool}

#     return context





# result = file_automation('../Review_n_Request/Preety_research paper_1.11 (1).docx', 1)
# print(result)